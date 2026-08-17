from pathlib import Path
import re
import sys

sys.path.insert(0, "/tmp/sisd232-ppt-deps")

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SCREENSHOTS = DOCS / "screenshots"
ASSETS = DOCS / "thesis-assets"
OUTPUT = DOCS / "skripsi-lengkap-awati-fujihani-lessy.docx"
LOGO = ROOT / "public" / "logo-malteng.png"

TITLE = "SISTEM INFORMASI JADWAL PELAJARAN DAN NILAI SISWA BERBASIS WEB PADA SD NEGERI 232 MALUKU TENGAH"
AUTHOR = "AWATI FUJIHANI LESSY"
STUDENT_ID = "220101161"
YEAR = "2026"


def font(size=12, bold=False, italic=False, color=None):
    return {
        "name": "Times New Roman",
        "size": Pt(size),
        "bold": bold,
        "italic": italic,
        "color": color,
    }


def apply_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Perbarui field di Microsoft Word"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_text, separate, text, end])
    return run


def add_page_number(section, fmt="decimal", start=1):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE")
    for run in paragraph.runs:
        apply_run_font(run, 10)
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    pg_num_type.set(qn("w:start"), str(start))


def set_update_fields(document):
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def load_font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_centered_text(draw, box, text, font_obj, fill, spacing=5):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font_obj, spacing=spacing, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - width) / 2, (y1 + y2 - height) / 2), text, font=font_obj, fill=fill, spacing=spacing, align="center")


def create_diagrams():
    ASSETS.mkdir(parents=True, exist_ok=True)
    title_font = load_font(42, True)
    body_font = load_font(28, False)
    small_font = load_font(23, False)

    waterfall = Image.new("RGB", (1900, 560), "white")
    draw = ImageDraw.Draw(waterfall)
    stages = [
        ("1", "Analisis\nKebutuhan"),
        ("2", "Perancangan\nSistem"),
        ("3", "Implementasi"),
        ("4", "Pengujian"),
        ("5", "Pemeliharaan"),
    ]
    colors = ["#0ea5e9", "#0284c7", "#0369a1", "#075985", "#0c4a6e"]
    for index, (number, label) in enumerate(stages):
        x = 65 + index * 365
        draw.rounded_rectangle((x, 150, x + 300, 410), radius=28, fill=colors[index], outline="#0f172a", width=2)
        draw.ellipse((x + 112, 95, x + 188, 171), fill="white", outline=colors[index], width=5)
        draw_centered_text(draw, (x + 112, 95, x + 188, 171), number, title_font, colors[index])
        draw_centered_text(draw, (x + 20, 180, x + 280, 380), label, body_font, "white")
        if index < len(stages) - 1:
            draw.line((x + 300, 280, x + 354, 280), fill="#64748b", width=8)
            draw.polygon([(x + 354, 260), (x + 354, 300), (x + 380, 280)], fill="#64748b")
    waterfall.save(ASSETS / "metode-waterfall.png")

    architecture = Image.new("RGB", (1900, 850), "white")
    draw = ImageDraw.Draw(architecture)
    draw.text((65, 45), "Arsitektur Sistem SISD 232", font=title_font, fill="#0f172a")
    roles = ["Admin", "Guru", "Siswa", "Kepala Sekolah"]
    for index, role in enumerate(roles):
        y = 150 + index * 155
        draw.rounded_rectangle((60, y, 330, y + 105), radius=20, fill="#e0f2fe", outline="#0284c7", width=4)
        draw_centered_text(draw, (60, y, 330, y + 105), role, body_font, "#075985")
        draw.line((330, y + 52, 500, 420), fill="#94a3b8", width=4)
    boxes = [
        ((500, 300, 850, 545), "Browser\nBlade + Livewire", "#f8fafc", "#475569"),
        ((950, 180, 1390, 400), "Laravel 13\nBusiness Logic\nAuthorization", "#e0f2fe", "#0369a1"),
        ((950, 510, 1390, 730), "Layanan Laporan\nDomPDF + Print", "#ecfdf5", "#15803d"),
        ((1510, 180, 1840, 400), "MySQL\nData Akademik", "#fff7ed", "#c2410c"),
        ((1510, 510, 1840, 730), "Storage\nLogo dan Aset", "#fdf2f8", "#be185d"),
    ]
    for coords, label, fill, outline in boxes:
        draw.rounded_rectangle(coords, radius=24, fill=fill, outline=outline, width=4)
        draw_centered_text(draw, coords, label, body_font, outline)
    arrows = [((850, 420), (950, 290)), ((850, 455), (950, 620)), ((1390, 290), (1510, 290)), ((1390, 620), (1510, 620))]
    for start, end in arrows:
        draw.line((*start, *end), fill="#64748b", width=6)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 22, ey - 12), (ex - 22, ey + 12)], fill="#64748b")
    architecture.save(ASSETS / "arsitektur-sistem.png")

    grade = Image.new("RGB", (1900, 600), "white")
    draw = ImageDraw.Draw(grade)
    steps = [
        "Guru memilih\npengajaran dan bulan",
        "Sistem memuat\nsiswa aktif",
        "Guru mengisi\nMinggu 1-4",
        "Validasi 0-100\ndan otorisasi",
        "Simpan massal\ndalam transaksi",
        "Rata-rata dan\nlaporan tersedia",
    ]
    for index, label in enumerate(steps):
        x = 40 + index * 310
        draw.rounded_rectangle((x, 175, x + 250, 430), radius=24, fill="#f8fafc", outline="#0284c7", width=4)
        draw.ellipse((x + 87, 110, x + 163, 186), fill="#0284c7")
        draw_centered_text(draw, (x + 87, 110, x + 163, 186), str(index + 1), body_font, "white")
        draw_centered_text(draw, (x + 18, 205, x + 232, 400), label, small_font, "#172033")
        if index < len(steps) - 1:
            draw.line((x + 250, 303, x + 292, 303), fill="#64748b", width=6)
            draw.polygon([(x + 292, 290), (x + 292, 316), (x + 310, 303)], fill="#64748b")
    grade.save(ASSETS / "alur-input-nilai.png")


create_diagrams()

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin = Cm(4)
section.right_margin = Cm(3)
section.header_distance = Cm(1.5)
section.footer_distance = Cm(1.5)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(0)

for style_name, size, bold, before, after in [
    ("Heading 1", 14, True, 0, 12),
    ("Heading 2", 12, True, 12, 6),
    ("Heading 3", 12, True, 9, 3),
]:
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

caption_style = styles["Caption"]
caption_style.font.name = "Times New Roman"
caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
caption_style.font.size = Pt(10)
caption_style.font.italic = False
caption_style.font.color.rgb = RGBColor(0, 0, 0)
caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption_style.paragraph_format.space_after = Pt(6)

if "Thesis Quote" not in styles:
    quote_style = styles.add_style("Thesis Quote", WD_STYLE_TYPE.PARAGRAPH)
else:
    quote_style = styles["Thesis Quote"]
quote_style.font.name = "Times New Roman"
quote_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
quote_style.font.size = Pt(11)
quote_style.font.italic = True
quote_style.paragraph_format.left_indent = Cm(1)
quote_style.paragraph_format.right_indent = Cm(1)
quote_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_center(text="", size=12, bold=False, spacing_after=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    apply_run_font(run, size, bold, italic)
    return p


def add_body(text, indent=True, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        apply_run_font(lead, 12, True)
        rest = p.add_run(text[len(bold_lead):])
        apply_run_font(rest)
    else:
        run = p.add_run(text)
        apply_run_font(run)
    return p


def add_bullets(items, numbered=False):
    for item in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        apply_run_font(run)


def add_chapter(number, title):
    doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"BAB {number}")
    apply_run_font(run, 14, True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    run2 = p2.add_run(title.upper())
    apply_run_font(run2, 14, True)


def add_heading(text, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    apply_run_font(run, 12 if level > 1 else 14, True)
    return p


def add_table(headers, rows, font_size=9, header_fill="DDEBF7", repeat=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header = table.rows[0]
    if repeat:
        set_repeat_table_header(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(label))
        apply_run_font(run, font_size, True)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            apply_run_font(run, font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(text):
    p = doc.add_paragraph(style="Caption")
    run = p.add_run(text)
    apply_run_font(run, 10)
    return p


def add_figure(path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(caption)


def add_placeholder(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    apply_run_font(run, 12, True, color=RGBColor(192, 0, 0))
    return p


# Cover
add_center("SKRIPSI", 14, True, 18)
add_center(TITLE, 14, True, 22)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(LOGO), height=Cm(4.6))
add_center("Diajukan sebagai salah satu syarat untuk memperoleh gelar Sarjana Komputer", 12, False, 16)
add_center("Disusun Oleh", 12, False, 4)
add_center(AUTHOR, 12, True, 2)
add_center(f"NIM: {STUDENT_ID}", 12, True, 20)
add_center("PROGRAM STUDI S1 TEKNIK INFORMATIKA", 12, True, 2)
add_placeholder("[NAMA FAKULTAS]")
add_placeholder("[NAMA UNIVERSITAS]")
add_placeholder(f"[KOTA] - {YEAR}")

# Front matter section
front_section = doc.add_section(WD_SECTION.NEW_PAGE)
front_section.page_width = Cm(21)
front_section.page_height = Cm(29.7)
front_section.top_margin = Cm(3)
front_section.bottom_margin = Cm(3)
front_section.left_margin = Cm(4)
front_section.right_margin = Cm(3)
add_page_number(front_section, "lowerRoman", 1)

add_center("LEMBAR PERSETUJUAN", 14, True, 22)
add_center(TITLE, 12, True, 16)
add_center(f"Oleh: {AUTHOR} / {STUDENT_ID}", 12, False, 22)
add_body("Skripsi ini telah diperiksa dan disetujui untuk diajukan dalam ujian skripsi Program Studi S1 Teknik Informatika.", indent=False)
add_center("[KOTA], [TANGGAL PERSETUJUAN]", 12, False, 20)
approval = add_table(["Pembimbing I", "Pembimbing II"], [["\n\n\n[NAMA PEMBIMBING I]\nNIDN: [NIDN]", "\n\n\n[NAMA PEMBIMBING II]\nNIDN: [NIDN]"]], 11)
add_center("Mengetahui,", 12, False, 8)
add_center("Ketua Program Studi S1 Teknik Informatika", 12, False, 32)
add_placeholder("[NAMA KETUA PROGRAM STUDI]")
add_center("NIDN: [NIDN]", 12)

doc.add_page_break()
add_center("LEMBAR PENGESAHAN", 14, True, 22)
add_center(TITLE, 12, True, 14)
add_center(f"Oleh: {AUTHOR} / {STUDENT_ID}", 12, False, 16)
add_body("Skripsi ini telah dipertahankan di hadapan Dewan Penguji dan dinyatakan memenuhi persyaratan akademik sesuai ketentuan perguruan tinggi.", indent=False)
add_center("Hari/Tanggal: [HARI, TANGGAL UJIAN]", 12, False, 10)
add_table(["Jabatan", "Nama", "Tanda Tangan"], [
    ["Ketua Penguji", "[NAMA KETUA PENGUJI]", "\n\n"],
    ["Penguji I", "[NAMA PENGUJI I]", "\n\n"],
    ["Penguji II", "[NAMA PENGUJI II]", "\n\n"],
], 11)
add_center("Mengetahui,", 12, False, 10)
add_placeholder("[DEKAN/KETUA PROGRAM STUDI SESUAI FORMAT KAMPUS]")

doc.add_page_break()
add_center("PERNYATAAN KEASLIAN", 14, True, 22)
for text in [
    f"Saya yang bertanda tangan di bawah ini, {AUTHOR}, NIM {STUDENT_ID}, menyatakan bahwa skripsi berjudul \"{TITLE.title()}\" merupakan hasil pekerjaan saya sendiri dengan arahan dosen pembimbing.",
    "Sumber informasi yang berasal atau dikutip dari karya pihak lain telah disebutkan dalam teks dan dicantumkan pada Daftar Pustaka. Apabila di kemudian hari ditemukan pelanggaran terhadap etika akademik, saya bersedia menerima sanksi sesuai peraturan yang berlaku.",
]:
    add_body(text)
add_center("[KOTA], [TANGGAL PERNYATAAN]", 12, False, 26)
add_center("Yang membuat pernyataan,", 12, False, 36)
add_center(AUTHOR, 12, True, 2)
add_center(f"NIM {STUDENT_ID}", 12)

doc.add_page_break()
add_center("ABSTRAK", 14, True, 16)
abstract_id = (
    "SD Negeri 232 Maluku Tengah memerlukan pengelolaan jadwal pelajaran dan nilai siswa yang terintegrasi untuk menggantikan proses manual yang berisiko menimbulkan keterlambatan pencarian data, kesalahan pencatatan, serta bentrok jadwal. Penelitian ini bertujuan merancang dan membangun sistem informasi berbasis web yang mendukung pengelolaan tahun ajaran, semester, guru, siswa, kelas, mata pelajaran, jam pelajaran, pengajaran, jadwal, nilai tugas mingguan, dan laporan. Metode pengembangan menggunakan Waterfall yang terdiri atas analisis kebutuhan, perancangan, implementasi, pengujian, dan pemeliharaan. Sistem dibangun menggunakan PHP 8.4, Laravel 13, Livewire 4, MySQL, Blade, Tailwind CSS, Spatie Laravel Permission, dan DomPDF. Empat jenis pengguna diterapkan, yaitu admin, guru, siswa, dan kepala sekolah. Validasi jadwal menolak bentrok kelas, bentrok guru, dan data duplikat. Nilai tugas disimpan untuk Minggu 1 sampai Minggu 4 setiap bulan dengan rentang 0-100, sedangkan nilai kosong dipertahankan sebagai NULL. Pengujian otomatis menghasilkan 18 pengujian lulus dengan 76 assertion, mencakup autentikasi, otorisasi, periode akademik, jadwal, nilai, laporan, dan pembatasan akses data. Hasil penelitian berupa aplikasi responsif yang menyediakan input nilai massal, monitoring, pencetakan, dan unduhan PDF. Sistem yang dibangun memenuhi kebutuhan fungsional utama dan menjaga integritas data akademik melalui foreign key, unique index, transaksi, validasi server, serta pembatasan akses berbasis role dan permission."
)
add_body(abstract_id)
p = doc.add_paragraph()
run = p.add_run("Kata kunci: ")
apply_run_font(run, 12, True)
run2 = p.add_run("sistem informasi; jadwal pelajaran; nilai siswa; Laravel; sekolah dasar")
apply_run_font(run2)

doc.add_page_break()
add_center("ABSTRACT", 14, True, 16)
abstract_en = (
    "SD Negeri 232 Maluku Tengah requires an integrated management system for lesson schedules and student grades to replace manual processes that may cause slow data retrieval, recording errors, and schedule conflicts. This study aims to design and develop a web-based information system for managing academic years, semesters, teachers, students, classes, subjects, lesson periods, teaching assignments, schedules, weekly assignment grades, and reports. The Waterfall method was employed through requirements analysis, system design, implementation, testing, and maintenance. The system was developed using PHP 8.4, Laravel 13, Livewire 4, MySQL, Blade, Tailwind CSS, Spatie Laravel Permission, and DomPDF. Four user roles were implemented: administrator, teacher, student, and principal. Schedule validation prevents class conflicts, teacher conflicts, and duplicate records. Assignment grades are stored for Week 1 to Week 4 of each month within a 0-100 range, while unassessed values remain NULL. Automated testing produced 18 passing tests with 76 assertions covering authentication, authorization, academic periods, schedules, grades, reports, and data access restrictions. The result is a responsive application that supports bulk grade entry, monitoring, printing, and PDF downloads. The system fulfills the main functional requirements and maintains academic data integrity through foreign keys, unique indexes, transactions, server-side validation, and role-based permissions."
)
add_body(abstract_en)
p = doc.add_paragraph()
run = p.add_run("Keywords: ")
apply_run_font(run, 12, True)
run2 = p.add_run("information system; lesson schedule; student grades; Laravel; elementary school")
apply_run_font(run2)

doc.add_page_break()
add_center("KATA PENGANTAR", 14, True, 16)
preface = [
    "Puji syukur penulis panjatkan ke hadirat Tuhan Yang Maha Esa karena atas penyertaan-Nya skripsi ini dapat disusun. Skripsi ini membahas perancangan dan pembangunan Sistem Informasi Jadwal Pelajaran dan Nilai Siswa Berbasis Web pada SD Negeri 232 Maluku Tengah.",
    "Penyusunan skripsi dan pengembangan aplikasi tidak terlepas dari dukungan berbagai pihak. Oleh karena itu, penulis menyampaikan terima kasih kepada pimpinan perguruan tinggi, dosen pembimbing, dosen penguji, Program Studi S1 Teknik Informatika, pihak SD Negeri 232 Maluku Tengah, keluarga, serta seluruh pihak yang memberikan dukungan selama proses penelitian.",
    "Penulis menyadari bahwa naskah ini masih dapat disempurnakan sesuai arahan pembimbing dan ketentuan institusi. Kritik dan saran yang membangun diharapkan dapat meningkatkan kualitas penelitian serta manfaat aplikasi bagi sekolah.",
]
for item in preface:
    add_body(item)
add_center("[KOTA], [TANGGAL]", 12, False, 18)
add_center("Penulis", 12, False, 26)
add_center(AUTHOR, 12, True)

doc.add_page_break()
add_center("DAFTAR ISI", 14, True, 16)
p = doc.add_paragraph()
add_field(p, 'TOC \\o "1-3" \\h \\z \\u')

doc.add_page_break()
add_center("DAFTAR GAMBAR", 14, True, 12)
figures = [
    "Gambar 2.1 Kerangka pikir penelitian",
    "Gambar 3.1 Tahapan metode Waterfall",
    "Gambar 3.2 Arsitektur sistem",
    "Gambar 3.3 Alur input nilai tugas",
    "Gambar 4.1 Halaman login",
    "Gambar 4.2 Dashboard Admin",
    "Gambar 4.3 Halaman pengajaran",
    "Gambar 4.4 Halaman jadwal pelajaran",
    "Gambar 4.5 Input nilai massal Guru",
    "Gambar 4.6 Halaman Nilai Saya",
    "Gambar 4.7 Laporan nilai",
    "Gambar 4.8 Tampilan aplikasi pada perangkat mobile",
]
for item in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(item)
    apply_run_font(run)

doc.add_page_break()
add_center("DAFTAR TABEL", 14, True, 12)
tables_list = [
    "Tabel 2.1 Penelitian terdahulu",
    "Tabel 3.1 Kebutuhan fungsional Admin",
    "Tabel 3.2 Kebutuhan fungsional Guru, Siswa, dan Kepala Sekolah",
    "Tabel 3.3 Kebutuhan nonfungsional",
    "Tabel 3.4 Entitas utama basis data",
    "Tabel 3.5 Jadwal penelitian",
    "Tabel 4.1 Lingkungan implementasi",
    "Tabel 4.2 Implementasi role dan permission",
    "Tabel 4.3 Hasil automated test",
    "Tabel 4.4 Ringkasan black box testing",
    "Tabel 4.5 Keterlacakan kebutuhan dan hasil",
]
for item in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(item)
    apply_run_font(run)

# Main section
main_section = doc.add_section(WD_SECTION.NEW_PAGE)
main_section.page_width = Cm(21)
main_section.page_height = Cm(29.7)
main_section.top_margin = Cm(3)
main_section.bottom_margin = Cm(3)
main_section.left_margin = Cm(4)
main_section.right_margin = Cm(3)
add_page_number(main_section, "decimal", 1)

add_chapter("I", "PENDAHULUAN")
add_heading("1.1 Latar Belakang", 2)
chapter1_background = [
    "Teknologi informasi menjadi salah satu sarana penting dalam mendukung pengelolaan data pada organisasi pendidikan. Sekolah tidak hanya membutuhkan media penyampaian informasi, tetapi juga memerlukan sistem yang mampu menyimpan, mengolah, menghubungkan, dan menyajikan data akademik secara konsisten. Sistem informasi berbasis web memberi keuntungan karena dapat diakses melalui peramban tanpa pemasangan aplikasi khusus pada setiap perangkat serta dapat digunakan dari komputer maupun telepon genggam.",
    "Pengelolaan akademik sekolah dasar mencakup berbagai data yang saling berkaitan, antara lain tahun ajaran, semester, guru, siswa, kelas, mata pelajaran, jam pelajaran, pembagian tugas mengajar, jadwal, dan nilai. Apabila data tersebut dikelola secara terpisah, proses pencarian dan penyusunan laporan menjadi lebih lambat. Perubahan pada satu data juga berisiko tidak segera tercermin pada catatan lain sehingga menurunkan konsistensi informasi.",
    "SD Negeri 232 Maluku Tengah masih menghadapi kebutuhan untuk meningkatkan pengelolaan jadwal pelajaran dan pencatatan nilai siswa yang sebelumnya dilakukan secara manual. Pada penyusunan jadwal, sekolah perlu memastikan bahwa satu kelas tidak menerima dua pelajaran pada waktu yang sama dan seorang guru tidak mengajar pada dua kelas secara bersamaan. Pemeriksaan secara manual membutuhkan ketelitian tinggi, terutama ketika jumlah kelas, mata pelajaran, dan alokasi jam bertambah.",
    "Pencatatan nilai tugas memiliki karakteristik khusus. Setiap guru mencatat nilai tugas siswa dari Minggu 1 sampai Minggu 4 pada setiap bulan. Nilai harus dapat dibedakan antara angka nol dan kondisi belum dinilai. Apabila kolom kosong diperlakukan sebagai nol, rata-rata siswa menjadi tidak tepat. Oleh sebab itu, sistem perlu menyimpan nilai kosong sebagai NULL dan menghitung rata-rata hanya dari minggu yang sudah memiliki nilai.",
    "Kebutuhan sekolah tidak berhenti pada pencatatan data. Guru membutuhkan antarmuka yang memungkinkan seluruh siswa pada satu kelas tampil dalam satu tabel agar input nilai tidak dilakukan melalui banyak halaman. Siswa membutuhkan akses untuk melihat jadwal kelas dan nilai miliknya sendiri. Kepala Sekolah membutuhkan informasi ringkas untuk monitoring, sedangkan Admin memerlukan fitur pengelolaan data serta laporan.",
    "Perbedaan kebutuhan tersebut menuntut penerapan hak akses yang tidak hanya berada pada tampilan. Pembatasan harus dilakukan pada route, proses bisnis, dan query data untuk mencegah pengguna mengakses data di luar kewenangannya. Guru tidak boleh mengubah nilai guru lain, siswa tidak boleh melihat nilai siswa lain, dan Kepala Sekolah tidak perlu memperoleh hak pengubahan penuh seperti Admin.",
    "Penelitian terdahulu menunjukkan bahwa sistem informasi akademik berbasis web dapat membantu pengelolaan data sekolah dan penyampaian informasi [1]-[5]. Meskipun demikian, setiap sekolah memiliki kebutuhan operasional yang berbeda. Penelitian ini berfokus pada integrasi jadwal pelajaran dengan nilai tugas mingguan dalam periode bulanan, histori penempatan kelas siswa, validasi bentrok jadwal, input nilai massal, dan laporan berbasis role pada konteks SD Negeri 232 Maluku Tengah.",
    "Sistem dikembangkan menggunakan Laravel karena menyediakan struktur aplikasi, autentikasi, validasi, ORM, migration, dan mekanisme keamanan yang mendukung pengembangan terpelihara [8]. Livewire digunakan untuk membangun antarmuka interaktif dengan tetap mempertahankan pendekatan server-driven [9]. MySQL digunakan sebagai basis data utama, sedangkan Spatie Laravel Permission menangani role dan permission. DomPDF digunakan untuk menghasilkan laporan PDF.",
    "Metode Waterfall dipilih karena tahapan penelitian pada proposal disusun berurutan, yaitu analisis kebutuhan, perancangan, implementasi, pengujian, dan pemeliharaan. Setiap tahap menghasilkan artefak yang menjadi masukan bagi tahap berikutnya. Pendekatan tersebut sesuai untuk proyek dengan lingkup yang telah didefinisikan secara rinci [6], [7].",
    "Berdasarkan kondisi tersebut, penelitian ini menghasilkan Sistem Informasi Jadwal Pelajaran dan Nilai Siswa Berbasis Web pada SD Negeri 232 Maluku Tengah. Sistem diharapkan membantu sekolah mengelola data secara terintegrasi, mencegah bentrok jadwal, mempercepat input nilai, memberikan akses informasi sesuai role, dan menghasilkan laporan yang dapat dicetak maupun diunduh sebagai PDF.",
]
for item in chapter1_background:
    add_body(item)

add_heading("1.2 Identifikasi Masalah", 2)
add_bullets([
    "Data jadwal, guru, kelas, mata pelajaran, dan nilai belum terintegrasi dalam satu sistem informasi.",
    "Pemeriksaan bentrok kelas dan bentrok guru pada jadwal membutuhkan ketelitian tinggi jika dilakukan secara manual.",
    "Pencatatan nilai tugas Minggu 1 sampai Minggu 4 belum memiliki mekanisme input massal dan perhitungan rata-rata yang membedakan nilai kosong dengan nilai nol.",
    "Informasi akademik perlu dibatasi sesuai kewenangan Admin, Guru, Siswa, dan Kepala Sekolah.",
    "Penyusunan laporan jadwal dan nilai membutuhkan proses yang lebih cepat, terstruktur, dapat dicetak, dan dapat diunduh sebagai PDF.",
])

add_heading("1.3 Rumusan Masalah", 2)
add_body("Bagaimana merancang dan membangun Sistem Informasi Jadwal Pelajaran dan Nilai Siswa Berbasis Web pada SD Negeri 232 Maluku Tengah yang mampu mengelola data akademik secara terintegrasi, mencegah bentrok jadwal, mendukung input nilai tugas Minggu 1 sampai Minggu 4, membatasi akses berdasarkan role, dan menghasilkan laporan yang dapat dicetak serta diunduh?", indent=False)

add_heading("1.4 Batasan Masalah", 2)
add_bullets([
    "Sistem digunakan untuk pengelolaan jadwal pelajaran dan nilai tugas mingguan, bukan untuk absensi, rapor lengkap, UTS, UAS, ujian daring, keuangan, atau penerimaan peserta didik baru.",
    "Nilai tugas disimpan untuk bulan 1 sampai 12 dan Minggu 1 sampai Minggu 4 dengan rentang 0 sampai 100 atau NULL jika belum dinilai.",
    "Pengguna sistem terdiri atas Admin, Guru, Siswa, dan Kepala Sekolah.",
    "Basis data utama adalah MySQL; SQLite hanya digunakan pada lingkungan automated test.",
    "Laporan mencakup laporan jadwal dan nilai berdasarkan filter yang tersedia serta dapat dicetak dan diunduh sebagai PDF.",
    "Sistem diimplementasikan sebagai aplikasi web responsif menggunakan Laravel, Livewire, Blade, Tailwind CSS, dan Vite.",
])

add_heading("1.5 Tujuan Penelitian", 2)
add_bullets([
    "Merancang basis data akademik yang menjaga hubungan dan histori data sekolah.",
    "Membangun aplikasi web untuk mengelola jadwal pelajaran dan nilai siswa secara terintegrasi.",
    "Menerapkan validasi bentrok kelas, bentrok guru, dan jadwal duplikat.",
    "Menerapkan input nilai massal Minggu 1 sampai Minggu 4 dengan perhitungan rata-rata yang tepat.",
    "Menerapkan autentikasi dan otorisasi berbasis role dan permission.",
    "Menyediakan monitoring, pencetakan, dan laporan PDF untuk pengguna yang berwenang.",
    "Menguji fungsi utama sistem melalui automated test dan black box testing.",
])

add_heading("1.6 Manfaat Penelitian", 2)
add_heading("1.6.1 Manfaat bagi Sekolah", 3)
add_body("Penelitian membantu sekolah memperoleh data akademik yang lebih terstruktur, mengurangi pengulangan pencatatan, mempercepat pencarian informasi, dan menyediakan laporan yang konsisten.")
add_heading("1.6.2 Manfaat bagi Guru", 3)
add_body("Guru dapat melihat jadwal mengajar dan mengisi nilai seluruh siswa dalam satu tabel untuk setiap pengajaran dan bulan, sehingga proses pencatatan lebih sederhana.")
add_heading("1.6.3 Manfaat bagi Siswa", 3)
add_body("Siswa memperoleh akses langsung terhadap jadwal kelas dan nilai miliknya sendiri tanpa dapat melihat data siswa lain.")
add_heading("1.6.4 Manfaat bagi Kepala Sekolah", 3)
add_body("Kepala Sekolah dapat memantau data, jadwal, nilai, dan laporan tanpa memperoleh akses pengubahan penuh terhadap master data.")
add_heading("1.6.5 Manfaat Akademik", 3)
add_body("Penelitian menjadi contoh penerapan Laravel dan Livewire pada sistem informasi akademik dengan aturan bisnis jadwal, histori kelas, input nilai mingguan, serta otorisasi berlapis.")

add_heading("1.7 Sistematika Penulisan", 2)
add_body("Bab I memuat latar belakang, masalah, batasan, tujuan, manfaat, dan sistematika penulisan. Bab II membahas teori pendukung, penelitian terdahulu, kebaruan, dan kerangka pikir. Bab III menjelaskan metode penelitian dan pengembangan, kebutuhan, rancangan, serta pengujian. Bab IV menyajikan hasil implementasi dan pembahasan. Bab V memuat kesimpulan dan saran. Bagian akhir berisi daftar pustaka dan lampiran.")

add_chapter("II", "TINJAUAN PUSTAKA DAN LANDASAN TEORI")
add_heading("2.1 Penelitian Terdahulu", 2)
add_body("Penelitian terdahulu digunakan untuk memahami pendekatan yang telah diterapkan pada sistem informasi akademik dan menentukan posisi penelitian ini. Lima penelitian pada proposal menjadi dasar perbandingan.")
add_caption("Tabel 2.1 Penelitian terdahulu")
add_table(["No", "Peneliti/Tahun", "Fokus", "Metode/Teknologi", "Perbedaan dengan Penelitian Ini"], [
    ["1", "Makkaraka dkk. (2024) [1]", "Sistem informasi akademik siswa berbasis web", "Perancangan aplikasi web", "Penelitian ini menambahkan validasi bentrok jadwal dan nilai tugas Minggu 1-4."],
    ["2", "Wahyudin dkk. (2023) [2]", "Sistem informasi akademik sekolah pada MTs", "Studi kasus berbasis web", "Objek penelitian berbeda dan sistem ini menerapkan empat role serta histori kelas."],
    ["3", "Lontaan dan Sinadia (2024) [3]", "Informasi akademik dan nonakademik sekolah", "Design and development", "Penelitian ini membatasi lingkup pada jadwal dan nilai dengan aturan bisnis lebih rinci."],
    ["4", "Prasetyo dkk. (2024) [4]", "Sistem informasi akademik website", "Extreme Programming", "Penelitian ini menggunakan Waterfall dan mengintegrasikan laporan PDF role-aware."],
    ["5", "Lengkong dkk. (2023) [5]", "Sistem informasi akademik sekolah kejuruan", "Aplikasi web", "Penelitian ini disesuaikan untuk sekolah dasar dan pencatatan nilai tugas bulanan."],
], 8)

add_heading("2.2 State of the Art dan Kebaruan", 2)
for item in [
    "State of the art penelitian berada pada integrasi pengelolaan jadwal dan nilai dalam satu aplikasi web. Sistem tidak hanya menampilkan informasi, tetapi mengimplementasikan aturan domain akademik melalui service layer, transaksi, unique index, foreign key, dan authorized query.",
    "Kebaruan kontekstual penelitian terletak pada pencatatan satu nilai untuk setiap kombinasi pengajaran, siswa, bulan, dan minggu. Model tersebut menjaga perbedaan antara nilai nol dan belum dinilai serta mendukung perhitungan rata-rata hanya dari nilai yang tersedia.",
    "Kebaruan fungsional meliputi histori kelas siswa lintas tahun ajaran, validasi bentrok kelas dan guru pada jadwal, input nilai massal, empat role, laporan role-aware, serta pembatasan siswa berdasarkan relasi akun login tanpa menerima siswa_id bebas dari browser.",
]:
    add_body(item)

theory_sections = [
    ("2.3 Sistem Informasi", [
        "Sistem informasi merupakan kesatuan komponen manusia, prosedur, data, perangkat lunak, dan teknologi yang bekerja untuk mengumpulkan, mengolah, menyimpan, dan menyajikan informasi. Dalam penelitian ini, komponen sistem informasi diwujudkan melalui pengguna sekolah, proses akademik, basis data, aplikasi web, dan laporan.",
        "Kualitas informasi dipengaruhi oleh ketepatan, kelengkapan, relevansi, dan ketersediaan. Penerapan constraint basis data dan validasi aplikasi bertujuan menjaga ketepatan, sedangkan filter dan laporan menyediakan informasi sesuai kebutuhan pengguna.",
    ]),
    ("2.4 Sistem Informasi Akademik", [
        "Sistem informasi akademik adalah aplikasi yang mendukung pengelolaan kegiatan pendidikan seperti data siswa, guru, kelas, mata pelajaran, jadwal, penilaian, dan laporan. Sistem akademik yang baik harus mampu mempertahankan histori karena data periode sebelumnya tetap diperlukan untuk pelacakan dan pelaporan.",
        "Pada sistem ini, histori diwujudkan melalui tahun ajaran, semester, pengajaran, jadwal, nilai, serta tabel siswa_kelas. Status aktif dan nonaktif digunakan untuk menghindari penghapusan data yang dapat merusak hubungan historis.",
    ]),
    ("2.5 Aplikasi Berbasis Web", [
        "Aplikasi berbasis web menggunakan protokol HTTP/HTTPS dan diakses melalui browser. Pendekatan ini memudahkan distribusi karena pengguna tidak perlu memasang aplikasi khusus. Pembaruan dilakukan pada server sehingga seluruh pengguna memperoleh versi yang sama.",
        "Aplikasi harus responsif agar dapat digunakan pada desktop dan telepon genggam. Antarmuka penelitian menggunakan pendekatan mobile-first, sidebar yang berubah menjadi drawer, tabel horizontal scroll, target sentuh yang memadai, dan formulir dengan label serta pesan validasi.",
    ]),
    ("2.6 Framework Laravel", [
        "Laravel adalah framework PHP yang menerapkan pola Model-View-Controller dan menyediakan routing, middleware, validation, Eloquent ORM, migration, authentication, authorization, storage, testing, dan fitur pendukung lain [8]. Laravel digunakan untuk membangun struktur backend dan proses bisnis aplikasi.",
        "Eloquent ORM menghubungkan model dengan tabel basis data. Migration mendefinisikan struktur tabel dan constraint secara terkontrol. Middleware melindungi route, sedangkan validation dan service class menangani aturan bisnis penting.",
    ]),
    ("2.7 Livewire, Blade, dan Tailwind CSS", [
        "Livewire memungkinkan antarmuka dinamis dibangun menggunakan komponen PHP yang berkomunikasi dengan server tanpa menulis banyak kode JavaScript [9]. Single-file component digunakan untuk menyatukan state, action, query, dan template halaman.",
        "Blade digunakan sebagai template engine, sedangkan Tailwind CSS digunakan untuk membentuk tampilan yang responsif dan konsisten. Alpine.js mendukung interaksi ringan seperti drawer mobile.",
    ]),
    ("2.8 Basis Data MySQL", [
        "MySQL adalah sistem manajemen basis data relasional. Data disimpan dalam tabel yang dihubungkan melalui primary key dan foreign key [10]. Sistem menggunakan unique index untuk mencegah duplikasi serta index pada kolom yang sering digunakan dalam pencarian dan filter.",
        "Integritas referensial dijaga menggunakan restrictOnDelete pada data akademik. Operasi penting seperti aktivasi periode, penempatan siswa, dan penyimpanan nilai massal menggunakan transaksi agar perubahan berhasil atau dibatalkan sebagai satu kesatuan.",
    ]),
    ("2.9 Role-Based Access Control", [
        "Role-Based Access Control memberikan izin berdasarkan peran pengguna. Implementasi menggunakan Spatie Laravel Permission [11] dengan role admin, guru, siswa, dan kepala_sekolah. Permission disimpan di basis data dan diterapkan pada route serta antarmuka.",
        "Pembatasan route belum cukup untuk mencegah akses objek yang tidak sah. Oleh sebab itu, sistem menambahkan validasi kepemilikan pengajaran dan query berbasis user login untuk mencegah insecure direct object reference sesuai prinsip kontrol akses OWASP [12].",
    ]),
    ("2.10 Jadwal Pelajaran", [
        "Jadwal pelajaran merepresentasikan pengajaran pada hari dan jam tertentu. Pengajaran menghubungkan semester, kelas, mata pelajaran, dan guru. Pemisahan ini mengurangi pengulangan data dan memudahkan validasi jadwal.",
        "Bentrok kelas terjadi apabila kelas yang sama memiliki lebih dari satu pengajaran pada semester, hari, dan jam yang sama. Bentrok guru terjadi apabila guru yang sama mengajar dua kelas pada waktu yang sama. Kedua kondisi diperiksa sebelum data tersimpan.",
    ]),
    ("2.11 Nilai Tugas Mingguan", [
        "Nilai tugas direkam per bulan dan minggu agar guru dapat memantau perkembangan tugas secara periodik. Rentang nilai adalah 0 sampai 100. Nilai NULL menunjukkan belum dinilai dan tidak disertakan dalam rata-rata.",
        "Rata-rata dihitung menggunakan jumlah nilai yang tersedia dibagi jumlah minggu yang sudah dinilai. Jika M1 bernilai 80 dan M2 bernilai 90 sedangkan M3 dan M4 kosong, rata-rata adalah 85, bukan 42,5.",
    ]),
    ("2.12 Metode Waterfall", [
        "Waterfall adalah pendekatan pengembangan berurutan yang mencakup analisis, desain, implementasi, pengujian, dan pemeliharaan [6], [7]. Kejelasan keluaran pada setiap tahap membantu pengendalian lingkup dan dokumentasi.",
        "Pada penelitian ini, requirement proposal menjadi dasar analisis. Rancangan data dan akses dibuat sebelum implementasi. Setelah fitur selesai, pengujian dilakukan secara bertahap dan hasilnya digunakan untuk perbaikan.",
    ]),
    ("2.13 Black Box Testing", [
        "Black box testing memeriksa kesesuaian keluaran terhadap input dan kebutuhan tanpa bergantung pada struktur internal program. Pengujian digunakan pada login, role, jadwal, nilai, laporan, dan tampilan mobile.",
        "Selain skenario black box, PHPUnit digunakan untuk menjalankan feature test secara otomatis. Pengujian otomatis memberikan hasil yang dapat diulang dan membantu mendeteksi regresi setelah perubahan kode.",
    ]),
]
for heading, paragraphs in theory_sections:
    add_heading(heading, 2)
    for paragraph in paragraphs:
        add_body(paragraph)

add_heading("2.14 Kerangka Pikir", 2)
add_body("Kerangka pikir dimulai dari permasalahan pengelolaan manual, dilanjutkan analisis kebutuhan, perancangan aplikasi terintegrasi, implementasi, pengujian, dan pencapaian hasil berupa pengelolaan akademik yang lebih terstruktur.")
add_figure(ASSETS / "metode-waterfall.png", "Gambar 2.1 Kerangka pikir penelitian", 6.0)

add_chapter("III", "METODOLOGI PENELITIAN")
add_heading("3.1 Jenis Penelitian", 2)
add_body("Penelitian ini merupakan penelitian rekayasa perangkat lunak terapan. Produk yang dihasilkan berupa aplikasi sistem informasi berbasis web. Proses penelitian berorientasi pada pemecahan masalah operasional sekolah melalui analisis kebutuhan, perancangan, pembangunan, dan pengujian sistem.")

add_heading("3.2 Lokasi dan Waktu Penelitian", 2)
add_body("Penelitian dilaksanakan pada SD Negeri 232 Maluku Tengah. Proposal merencanakan pelaksanaan selama dua bulan, yaitu Mei sampai Juni, meliputi analisis kebutuhan, perancangan, implementasi, pengujian, pemeliharaan awal, dan penyusunan laporan. Tanggal rinci perlu disesuaikan dengan catatan pelaksanaan dan kalender akademik kampus.")

add_heading("3.3 Objek dan Subjek Penelitian", 2)
add_body("Objek penelitian adalah proses pengelolaan jadwal pelajaran dan nilai tugas siswa. Subjek yang terkait dengan kebutuhan sistem meliputi pengelola sekolah, guru, siswa, dan kepala sekolah. Dokumen ini tidak mengklaim hasil wawancara atau kuesioner yang belum dilaksanakan; bukti pengumpulan data lapangan perlu dilampirkan apabila menjadi persyaratan kampus.")

add_heading("3.4 Jenis dan Sumber Data", 2)
add_bullets([
    "Data primer berupa kebutuhan proses sekolah yang dirumuskan pada proposal dan klarifikasi kebutuhan dengan pihak sekolah.",
    "Data sekunder berupa literatur sistem informasi akademik, dokumentasi framework, standar keamanan, dan dokumen teknis aplikasi.",
    "Data pengujian berupa data demo yang konsisten secara relasional untuk menguji guru, siswa, kelas, jadwal, dan nilai.",
])

add_heading("3.5 Teknik Pengumpulan Data", 2)
add_heading("3.5.1 Observasi Proses", 3)
add_body("Observasi diarahkan pada alur pengelolaan jadwal, pembagian guru mengajar, pencatatan nilai mingguan, dan kebutuhan laporan. Hasil observasi dituangkan menjadi kebutuhan fungsional dan aturan bisnis.")
add_heading("3.5.2 Wawancara/Klarifikasi Kebutuhan", 3)
add_body("Klarifikasi dilakukan kepada pihak sekolah untuk memastikan istilah, pengguna, periode nilai, dan keluaran laporan. Transkrip atau berita acara perlu dilampirkan berdasarkan kegiatan lapangan yang benar-benar dilaksanakan.")
add_heading("3.5.3 Studi Pustaka", 3)
add_body("Studi pustaka dilakukan terhadap penelitian terdahulu, konsep rekayasa perangkat lunak, dokumentasi Laravel, Livewire, MySQL, role-permission, keamanan aplikasi, dan pengujian.")

add_heading("3.6 Metode Pengembangan Waterfall", 2)
add_figure(ASSETS / "metode-waterfall.png", "Gambar 3.1 Tahapan metode Waterfall", 6.0)
waterfall_steps = [
    ("3.6.1 Analisis Kebutuhan", "Tahap ini mengidentifikasi data, pengguna, fungsi, aturan bisnis, batasan, dan keluaran. Luaran berupa requirement, role-permission, alur sistem, serta prioritas implementasi."),
    ("3.6.2 Perancangan Sistem", "Perancangan mencakup arsitektur, basis data relasional, navigasi, antarmuka, validasi, dan laporan. Luaran berupa ERD, relasi model, route, service layer, dan rancangan halaman."),
    ("3.6.3 Implementasi", "Tahap implementasi menerjemahkan rancangan menjadi kode Laravel dan Livewire, migration MySQL, model Eloquent, service, middleware, halaman responsif, laporan PDF, seeder, dan factory."),
    ("3.6.4 Pengujian", "Pengujian meliputi feature test, black box testing, migration bersih, seeder, build frontend, formatter, dan audit dependency. Temuan pengujian diperbaiki sebelum quality gate akhir."),
    ("3.6.5 Pemeliharaan", "Pemeliharaan mencakup perbaikan kesalahan, pembaruan dependency yang kompatibel, backup, monitoring log, penyesuaian kebutuhan sekolah, dan pengujian regresi."),
]
for heading, paragraph in waterfall_steps:
    add_heading(heading, 3)
    add_body(paragraph)

add_heading("3.7 Analisis Kebutuhan Fungsional", 2)
add_caption("Tabel 3.1 Kebutuhan fungsional Admin")
add_table(["Kode", "Kebutuhan", "Keluaran"], [
    ["ADM-01", "Login, logout, dan dashboard", "Akses aman dan ringkasan data"],
    ["ADM-02", "Kelola pengguna, guru, siswa, kelas, mata pelajaran, jam, periode", "CRUD master data"],
    ["ADM-03", "Penempatan siswa ke kelas", "Histori kelas per tahun ajaran"],
    ["ADM-04", "Kelola pengajaran dan jadwal", "Jadwal tanpa bentrok"],
    ["ADM-05", "Monitoring nilai", "Data nilai seluruh periode"],
    ["ADM-06", "Pengaturan sekolah", "Identitas dan logo laporan"],
    ["ADM-07", "Laporan, print, dan PDF", "Dokumen jadwal dan nilai"],
], 9)
add_caption("Tabel 3.2 Kebutuhan fungsional Guru, Siswa, dan Kepala Sekolah")
add_table(["Role", "Kebutuhan Utama", "Batasan"], [
    ["Guru", "Dashboard, jadwal sendiri, input/update nilai massal, riwayat, laporan", "Hanya pengajaran milik guru"],
    ["Siswa", "Dashboard, jadwal kelas, nilai sendiri, laporan", "Tidak dapat melihat siswa lain atau mengubah data"],
    ["Kepala Sekolah", "Dashboard monitoring, jadwal, nilai, laporan", "Akses baca, tidak memperoleh CRUD penuh"],
], 9)

add_heading("3.8 Kebutuhan Nonfungsional", 2)
add_caption("Tabel 3.3 Kebutuhan nonfungsional")
add_table(["Aspek", "Kebutuhan"], [
    ["Keamanan", "CSRF, autentikasi, rate limit login, role-permission, validasi server, hashing password, upload aman, pencegahan IDOR."],
    ["Kinerja", "Pagination, eager loading, index, query berbasis filter, transaksi dan upsert nilai."],
    ["Kegunaan", "Bahasa Indonesia, label jelas, loading, empty state, notifikasi, konfirmasi, responsif."],
    ["Kompatibilitas", "Browser modern pada desktop, tablet, dan mobile; MySQL untuk production."],
    ["Pemeliharaan", "Migration, seeder, service class, automated test, dokumentasi instalasi dan deployment."],
], 9)

add_heading("3.9 Arsitektur Sistem", 2)
add_body("Arsitektur menggunakan pola aplikasi web monolitik terstruktur. Browser menampilkan Blade dan Livewire. Request diproses Laravel melalui middleware, komponen, controller, service, model Eloquent, dan basis data MySQL. Laporan PDF diproses DomPDF, sedangkan berkas logo disimpan melalui Laravel Storage.")
add_figure(ASSETS / "arsitektur-sistem.png", "Gambar 3.2 Arsitektur sistem", 6.1)

add_heading("3.10 Perancangan Basis Data", 2)
add_body("Basis data memisahkan master, periode, transaksi akademik, dan otorisasi. Tabel siswa_kelas digunakan sebagai sumber histori penempatan sehingga kelas tidak disimpan sebagai satu-satunya atribut langsung pada siswa.")
add_caption("Tabel 3.4 Entitas utama basis data")
entity_rows = [
    ["users", "Akun, username, password, status, last login"],
    ["pengaturan_sekolah", "Identitas sekolah, logo, kepala sekolah"],
    ["tahun_ajaran, semester", "Periode akademik dan status aktif"],
    ["guru, siswa", "Profil akademik dan relasi akun"],
    ["kelas, siswa_kelas", "Kelas dan histori penempatan siswa"],
    ["mata_pelajaran, jam_pelajaran", "Referensi mata pelajaran dan alokasi waktu"],
    ["pengajaran", "Relasi semester, kelas, mata pelajaran, dan guru"],
    ["jadwal_pelajaran", "Hari dan jam untuk suatu pengajaran"],
    ["nilai_tugas", "Nilai per pengajaran, siswa, bulan, dan minggu"],
    ["roles, permissions", "Hak akses pengguna"],
]
add_table(["Entitas", "Fungsi"], entity_rows, 9)

add_heading("3.11 Perancangan Proses Nilai", 2)
add_body("Proses nilai dirancang untuk menghindari query per siswa. Sistem mengambil siswa aktif dan nilai periode terpilih secara batch. Pada penyimpanan, service memvalidasi guru, pengajaran, keanggotaan kelas, bulan, minggu, dan rentang nilai kemudian melakukan upsert dalam transaksi.")
add_figure(ASSETS / "alur-input-nilai.png", "Gambar 3.3 Alur input nilai tugas", 6.1)

add_heading("3.12 Perancangan Pengujian", 2)
add_body("Pengujian menggabungkan automated feature test dan skenario black box. Automated test menggunakan database terisolasi agar dapat diulang. Black box testing memeriksa input dan keluaran dari sudut pandang pengguna. Quality gate mencakup migrate fresh dan seed, test suite, Blade compilation, formatter, Composer validation, dependency audit, route audit, dan npm build.")

add_heading("3.13 Indikator Keberhasilan", 2)
add_bullets([
    "Seluruh role dapat login dan hanya mengakses menu yang diizinkan.",
    "Data master dan transaksi akademik dapat dikelola sesuai aturan.",
    "Bentrok kelas, bentrok guru, dan jadwal duplikat ditolak.",
    "Nilai 0, 100, dan NULL diterima; nilai di luar rentang ditolak.",
    "Guru lain dan siswa di luar kelas tidak dapat memanipulasi nilai.",
    "Laporan web, print, dan PDF dapat dihasilkan.",
    "Automated test dan build production selesai tanpa kegagalan.",
])

add_heading("3.14 Jadwal Penelitian", 2)
add_caption("Tabel 3.5 Jadwal penelitian")
add_table(["Kegiatan", "Mei Minggu 1", "Mei Minggu 2", "Mei Minggu 3", "Mei Minggu 4", "Juni Minggu 1-2", "Juni Minggu 3-4"], [
    ["Analisis kebutuhan", "Ya", "Ya", "", "", "", ""],
    ["Perancangan", "", "Ya", "Ya", "", "", ""],
    ["Implementasi", "", "", "Ya", "Ya", "Ya", ""],
    ["Pengujian dan perbaikan", "", "", "", "", "Ya", "Ya"],
    ["Dokumentasi/laporan", "", "", "", "", "", "Ya"],
], 7.5)

add_chapter("IV", "HASIL DAN PEMBAHASAN")
add_heading("4.1 Hasil Implementasi", 2)
add_body("Hasil penelitian adalah aplikasi SISD 232 yang dapat dijalankan sebagai aplikasi web dengan empat role. Implementasi mengikuti kebutuhan pada proposal dan diperluas dengan integritas basis data, histori kelas, validasi jadwal, input nilai massal, monitoring, laporan, dokumentasi, dan automated test.")

add_heading("4.2 Lingkungan Implementasi", 2)
add_caption("Tabel 4.1 Lingkungan implementasi")
add_table(["Komponen", "Versi/Fungsi"], [
    ["PHP", "8.4.20"],
    ["Laravel", "13.25 pada audit implementasi"],
    ["Livewire", "4.4 single-file components"],
    ["MySQL", "Basis data development/production"],
    ["Tailwind CSS", "4.x"],
    ["Vite", "8.x"],
    ["Spatie Permission", "8.3"],
    ["DomPDF", "3.x"],
    ["PHPUnit", "12.x"],
], 9)

add_heading("4.3 Implementasi Autentikasi", 2)
add_body("Halaman login menerima username atau email. Login diberi rate limiting dan menolak akun nonaktif. Registrasi publik tidak tersedia. Setelah login, pengguna diarahkan ke dashboard dan menu disusun sesuai role.")
add_figure(SCREENSHOTS / "00-login.png", "Gambar 4.1 Halaman login", 5.9)

add_heading("4.4 Dashboard Role-Aware", 2)
add_body("Dashboard Admin dan Kepala Sekolah menampilkan ringkasan jumlah guru, siswa, kelas, mata pelajaran, jadwal, dan nilai. Dashboard Guru memprioritaskan jadwal hari ini dan jalan cepat input nilai. Dashboard Siswa menampilkan kelas, jadwal, dan ringkasan nilai sendiri.")
add_figure(SCREENSHOTS / "admin-dashboard.png", "Gambar 4.2 Dashboard Admin", 5.9)

add_heading("4.5 Implementasi Master Data dan Pengajaran", 2)
add_body("Halaman master data menggunakan pola tabel yang konsisten dengan pencarian, filter status, sorting, pilihan jumlah data, pagination, empty state, modal form, validasi, loading, dan konfirmasi tindakan. Pengajaran menyimpan hubungan semester, kelas, mata pelajaran, dan guru tanpa menduplikasi atribut tersebut pada jadwal dan nilai.")
add_figure(SCREENSHOTS / "admin-pengajaran.png", "Gambar 4.3 Halaman pengajaran", 5.9)

add_heading("4.6 Implementasi Jadwal Pelajaran", 2)
add_body("Jadwal mengacu pada pengajaran, hari, dan jam pelajaran. Sebelum menyimpan, ValidasiJadwal memeriksa kelas dan guru pada semester, hari, dan jam yang sama. Unique constraint mencegah record identik. Guru hanya melihat jadwal miliknya dan siswa hanya melihat jadwal kelas aktif.")
add_figure(SCREENSHOTS / "admin-jadwal-pelajaran.png", "Gambar 4.4 Halaman jadwal pelajaran", 5.9)

add_heading("4.7 Implementasi Nilai Tugas", 2)
add_body("Guru memilih pengajaran dan bulan. Sistem menampilkan seluruh siswa aktif pada kelas tersebut beserta empat input minggu dan rata-rata. Penyimpanan dilakukan massal melalui PenyimpananNilaiMassal. Guru hanya dapat menyimpan pengajaran yang terkait dengan profilnya.")
add_figure(SCREENSHOTS / "guru-input-nilai-terisi.png", "Gambar 4.5 Input nilai massal Guru", 5.9)
add_body("Perhitungan rata-rata menggunakan koleksi nilai yang tidak kosong. Pendekatan ini memenuhi aturan bahwa NULL berarti belum dinilai, sedangkan angka nol adalah nilai yang sah. Unique key pada nilai_tugas memastikan hanya satu nilai untuk kombinasi pengajaran, siswa, bulan, dan minggu.")

add_heading("4.8 Implementasi Akses Siswa", 2)
add_body("Halaman Nilai Saya memperoleh siswa melalui user yang sedang login. Query tidak menerima siswa_id bebas dari browser. Nilai dikelompokkan per pengajaran dan ditampilkan bersama mata pelajaran, guru, M1-M4, dan rata-rata.")
add_figure(SCREENSHOTS / "siswa-nilai-saya.png", "Gambar 4.6 Halaman Nilai Saya", 5.9)

add_heading("4.9 Implementasi Laporan", 2)
add_body("Laporan jadwal dan nilai menyediakan filter periode dan data akademik. Print stylesheet menyembunyikan sidebar, tombol, filter, dan pagination. Endpoint PDF mengulangi pembatasan Guru dan Siswa agar parameter URL tidak memperluas akses.")
add_figure(SCREENSHOTS / "admin-laporan-nilai.png", "Gambar 4.7 Laporan nilai", 5.9)

add_heading("4.10 Implementasi Responsif", 2)
add_body("Pada perangkat mobile, sidebar berubah menjadi drawer. Form menyesuaikan satu kolom dan tabel lebar menggunakan horizontal scroll. Halaman nilai mempertahankan nama siswa sebagai kolom sticky sehingga guru tetap mengetahui baris yang sedang diisi.")
add_figure(SCREENSHOTS / "mobile-guru-input-nilai-terisi.png", "Gambar 4.8 Tampilan aplikasi pada perangkat mobile", 2.7)

add_heading("4.11 Implementasi Role dan Permission", 2)
add_caption("Tabel 4.2 Implementasi role dan permission")
add_table(["Area", "Admin", "Guru", "Siswa", "Kepala Sekolah"], [
    ["Dashboard", "Penuh", "Data sendiri", "Data sendiri", "Monitoring"],
    ["Master data", "CRUD", "Baca terbatas", "Tidak", "Baca terbatas"],
    ["Jadwal", "CRUD", "Milik guru", "Kelas aktif", "Seluruh data"],
    ["Nilai", "Monitoring", "Input/update sendiri", "Nilai sendiri", "Monitoring"],
    ["Laporan", "Seluruh data", "Pengajaran sendiri", "Data sendiri", "Seluruh data"],
    ["Pengaturan", "Update", "Tidak", "Tidak", "Tidak"],
], 8)

add_heading("4.12 Integritas dan Keamanan", 2)
security_paragraphs = [
    "CSRF token digunakan pada form, password disimpan menggunakan hashed cast, output Blade di-escape secara default, dan upload logo divalidasi berdasarkan tipe serta ukuran. Konfigurasi sensitif disimpan pada environment variable.",
    "Foreign key menjaga hubungan data dan restrict deletion melindungi histori. Data guru, siswa, dan pengguna yang sudah digunakan dinonaktifkan daripada dihapus. Unique index digunakan pada username, identitas opsional, periode, pengajaran, jadwal, dan nilai.",
    "Aktivasi tahun ajaran dan semester berjalan dalam transaksi. Penempatan siswa memastikan satu kelas aktif per tahun ajaran. Nilai massal menggunakan transaksi dan upsert untuk mengurangi query serta menjaga konsistensi.",
    "Otorisasi diterapkan melalui middleware permission/role, service validation, dan query yang berangkat dari user login. Pendekatan berlapis ini mengurangi risiko broken access control dan IDOR.",
]
for item in security_paragraphs:
    add_body(item)

add_heading("4.13 Hasil Automated Test", 2)
add_body("Test suite dijalankan setelah migration dan seeder. Hasil akhir menunjukkan 18 test lulus dengan 76 assertion. Pengujian mencakup autentikasi, role, konflik jadwal, nilai, periode, duplikasi, PDF, rendering halaman, dan validasi Livewire.")
add_caption("Tabel 4.3 Hasil automated test")
add_table(["Kelompok", "Skenario Utama", "Hasil"], [
    ["Autentikasi", "Login, logout, password salah, user nonaktif, register 404", "Lulus"],
    ["Otorisasi", "Siswa/Guru/Kepala Sekolah ditolak dari halaman Admin", "Lulus"],
    ["Jadwal", "Bentrok kelas dan guru", "Lulus"],
    ["Nilai", "0, 100, NULL, rentang, minggu, siswa luar kelas, guru lain", "Lulus"],
    ["Periode", "Satu tahun aktif, semester pada tahun aktif", "Lulus"],
    ["Data relasional", "Duplikasi pengajaran dan edit foreign key", "Lulus"],
    ["Laporan", "Endpoint PDF jadwal dan nilai", "Lulus"],
    ["UI runtime", "Halaman Admin utama dapat dirender", "Lulus"],
], 9)

add_heading("4.14 Hasil Black Box Testing", 2)
add_body("Dokumen black box memuat 22 skenario. Enam belas skenario telah terwakili oleh automated test, sedangkan enam skenario visual/interaksi disiapkan untuk pengujian manual. Status manual tidak dinyatakan lulus sebelum benar-benar diuji oleh pengguna.")
add_caption("Tabel 4.4 Ringkasan black box testing")
add_table(["Status", "Jumlah", "Keterangan"], [
    ["Lulus otomatis", "16", "Diverifikasi melalui test suite"],
    ["Siap uji manual", "6", "Guru CRUD, identitas duplikat, penempatan, filter, print, dan mobile"],
    ["Total", "22", "Seluruh skenario terdokumentasi"],
], 9)

add_heading("4.15 Quality Gate", 2)
add_bullets([
    "Migration fresh dan seeder berhasil pada database pengujian terisolasi.",
    "PHPUnit: 18 test lulus dan 76 assertion.",
    "Blade view compilation berhasil.",
    "Laravel Pint berhasil tanpa pelanggaran format.",
    "npm run build berhasil menghasilkan aset production.",
    "Composer validation berhasil; audit dependency production tidak menemukan advisory pada saat pemeriksaan.",
    "Storage link, route, screenshot role, PDF, dan dokumentasi telah diverifikasi.",
])

add_heading("4.16 Keterlacakan Kebutuhan", 2)
add_caption("Tabel 4.5 Keterlacakan kebutuhan dan hasil")
add_table(["Kebutuhan", "Implementasi", "Bukti"], [
    ["Empat role", "Spatie Permission dan middleware", "Dashboard dan authorization test"],
    ["Jadwal bebas bentrok", "ValidasiJadwal dan unique constraint", "Test bentrok kelas/guru"],
    ["Nilai Minggu 1-4", "nilai_tugas dan input massal", "Test nilai 0/100/NULL"],
    ["Siswa melihat nilai sendiri", "Authorized query dari user login", "Test halaman Nilai Saya"],
    ["Laporan PDF", "DomPDF dan filter role-aware", "Test content-type PDF"],
    ["Responsif", "Tailwind, drawer, horizontal scroll", "Screenshot 390px"],
    ["Histori kelas", "siswa_kelas dan status", "Service penempatan"],
], 8)

add_heading("4.17 Pembahasan", 2)
discussion = [
    "Hasil implementasi menunjukkan bahwa integrasi data mengurangi kebutuhan memasukkan atribut guru, kelas, dan mata pelajaran berulang kali pada setiap jadwal atau nilai. Pengajaran menjadi pusat hubungan dan memudahkan pembatasan akses guru.",
    "Validasi jadwal menyelesaikan risiko utama pada proses manual. Pemeriksaan dilakukan terhadap kelas dan guru dalam semester yang sama. Database constraint menambah perlindungan terhadap duplikasi identik, sedangkan service menangani konflik lintas pengajaran.",
    "Desain nilai per minggu memberikan fleksibilitas untuk nilai kosong dan nol. Input massal lebih sesuai dengan pola kerja guru dibandingkan form per siswa. Penggunaan upsert dan transaction mendukung efisiensi dan konsistensi saat jumlah siswa bertambah.",
    "Penerapan empat role memenuhi kebutuhan pengguna yang berbeda. Pembatasan server-side penting karena penyembunyian menu tidak mencegah manipulasi URL. Pengujian menunjukkan pengguna tanpa izin menerima respons yang sesuai.",
    "Dibandingkan penelitian terdahulu [1]-[5], penelitian ini menekankan kombinasi jadwal bebas bentrok, histori kelas, nilai mingguan bulanan, dan laporan role-aware pada sekolah dasar. Kontribusi utama bersifat kontekstual dan rekayasa, yaitu penerapan aturan bisnis sekolah ke dalam aplikasi yang dapat diuji.",
    "Pengujian otomatis memberi bukti kuat pada aturan kritis, tetapi tidak menggantikan evaluasi penerimaan pengguna. Uji manual bersama pihak sekolah, kuesioner usability, dan pengamatan penggunaan produksi tetap diperlukan untuk mengukur kepuasan, kemudahan belajar, dan dampak operasional.",
]
for item in discussion:
    add_body(item)

add_heading("4.18 Keterbatasan Penelitian", 2)
add_bullets([
    "Pengujian manual bersama pengguna sekolah belum didokumentasikan sebagai hasil lulus; dokumen menyediakan skenario yang siap digunakan.",
    "Pengujian quality gate database dilakukan pada lingkungan terisolasi yang kompatibel; migration tetap perlu dijalankan pada server MySQL target sebelum deployment.",
    "Sistem tidak mencakup absensi, rapor lengkap, UTS, UAS, keuangan, PPDB, atau pembelajaran daring.",
    "Evaluasi kuantitatif usability dan perbandingan waktu kerja sebelum-sesudah belum tersedia sehingga tidak boleh disimpulkan tanpa data lapangan.",
])

add_chapter("V", "PENUTUP")
add_heading("5.1 Kesimpulan", 2)
conclusions = [
    "Sistem Informasi Jadwal Pelajaran dan Nilai Siswa Berbasis Web pada SD Negeri 232 Maluku Tengah berhasil dirancang dan dibangun menggunakan metode Waterfall serta teknologi Laravel, Livewire, MySQL, Blade, Tailwind CSS, Spatie Permission, dan DomPDF.",
    "Sistem mengintegrasikan pengguna, guru, siswa, histori kelas, mata pelajaran, jam, tahun ajaran, semester, pengajaran, jadwal, nilai tugas Minggu 1-4, pengaturan sekolah, dan laporan.",
    "Aturan jadwal dapat menolak bentrok kelas, bentrok guru, dan jadwal duplikat. Nilai menerima rentang 0-100 dan NULL, sedangkan rata-rata hanya menghitung nilai yang tersedia.",
    "Role dan permission membatasi akses Admin, Guru, Siswa, dan Kepala Sekolah. Guru dibatasi pada pengajarannya dan siswa dibatasi pada data sendiri melalui query server-side.",
    "Aplikasi menyediakan dashboard, input nilai massal, monitoring, print, PDF, tampilan responsif, serta dokumentasi penggunaan dan deployment.",
    "Hasil quality gate menunjukkan 18 automated test dengan 76 assertion lulus dan build production berhasil. Dengan demikian, kebutuhan fungsional kritis telah diimplementasikan dan diverifikasi pada lingkungan pengembangan/pengujian.",
]
add_bullets(conclusions, numbered=True)

add_heading("5.2 Saran", 2)
add_bullets([
    "Melaksanakan User Acceptance Testing bersama perwakilan Admin, Guru, Siswa, dan Kepala Sekolah serta menyimpan berita acara hasil pengujian.",
    "Melakukan evaluasi usability menggunakan instrumen yang tervalidasi setelah pengguna memiliki pengalaman menggunakan aplikasi.",
    "Menyiapkan backup database terjadwal, HTTPS, monitoring log, dan prosedur pemulihan pada deployment production.",
    "Melakukan pengujian beban apabila jumlah pengguna atau data meningkat signifikan.",
    "Menambahkan fitur impor data hanya setelah proses inti stabil dan kebutuhan sekolah telah disetujui, tanpa memperluas sistem ke fitur di luar lingkup penelitian.",
])

doc.add_page_break()
add_center("DAFTAR PUSTAKA", 14, True, 16)
references = [
    "[1] A. M. R. B. Makkaraka, A. Iskandar, and W. Yang, \"Design of web-based student academic information system,\" Ceddi Journal of Education, vol. 3, no. 2, pp. 9-15, 2024.",
    "[2] W. Wahyudin, H. Wahyudi, and K. Komarudin, \"Web-Based School Academic Information System: Case Study at an MTs School in Bandung,\" Majalah Bisnis & IPTEK, vol. 16, no. 1, pp. 26-34, 2023.",
    "[3] R. J. Lontaan and A. R. Sinadia, \"Design and development of a web-based school information system,\" CogITo Smart Journal, vol. 10, no. 2, pp. 593-606, 2024.",
    "[4] D. Prasetyo, A. Utami, and T. G. Laksana, \"Website based academic information system design using Extreme Programming method,\" Journal of Informatics Information System Software Engineering and Applications, vol. 6, no. 2, pp. 134-143, 2024.",
    "[5] J. S. J. Lengkong, S. N. H. Jacobus, R. Dondokambey, K. F. Ratumbuisang, D. Paath, and E. S. Liow, \"Web-Based Academic Information Systems in Vocational School,\" International Journal of Information Technology and Education, vol. 2, no. 4, pp. 12-25, 2023.",
    "[6] R. S. Pressman and B. R. Maxim, Software Engineering: A Practitioner's Approach, 9th ed. New York: McGraw-Hill Education, 2019.",
    "[7] I. Sommerville, Software Engineering, 10th ed. Boston: Pearson, 2015.",
    "[8] Laravel, \"Laravel Documentation,\" Laravel Holdings Inc. [Online]. Available: https://laravel.com/docs. [Accessed: 18-Aug-2026].",
    "[9] Livewire, \"Livewire Documentation,\" [Online]. Available: https://livewire.laravel.com/docs. [Accessed: 18-Aug-2026].",
    "[10] Oracle, \"MySQL Reference Manual,\" [Online]. Available: https://dev.mysql.com/doc/. [Accessed: 18-Aug-2026].",
    "[11] Spatie, \"Laravel Permission Documentation,\" [Online]. Available: https://spatie.be/docs/laravel-permission. [Accessed: 18-Aug-2026].",
    "[12] OWASP Foundation, \"OWASP Application Security Verification Standard,\" [Online]. Available: https://owasp.org/www-project-application-security-verification-standard/. [Accessed: 18-Aug-2026].",
    "[13] ISO/IEC, ISO/IEC 25010:2023 Systems and software engineering - Systems and software Quality Requirements and Evaluation (SQuaRE) - Product quality model. Geneva: ISO, 2023.",
    "[14] ISTQB, \"Certified Tester Foundation Level Syllabus,\" International Software Testing Qualifications Board, 2024.",
]
for reference in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(reference)
    apply_run_font(run, 11)

doc.add_page_break()
add_center("LAMPIRAN", 14, True, 16)
add_heading("Lampiran 1. Matriks Black Box Testing", 2)
black_box_rows = [
    ["1", "Login", "Kredensial benar", "Masuk dashboard", "Lulus otomatis"],
    ["2", "Login", "Password salah", "Ditolak dengan pesan", "Lulus otomatis"],
    ["3", "Login", "User nonaktif", "Login ditolak", "Lulus otomatis"],
    ["4", "Registrasi", "Akses /register", "404", "Lulus otomatis"],
    ["5", "Role", "Siswa membuka pengguna", "403", "Lulus otomatis"],
    ["6", "Role", "Guru membuka tahun ajaran", "403", "Lulus otomatis"],
    ["7", "Guru", "Tambah/ubah data valid", "Data tersimpan", "Siap uji manual"],
    ["8", "Siswa", "NIS/NISN duplikat", "Validasi gagal", "Siap uji manual"],
    ["9", "Kelas", "Penempatan bulk", "Satu kelas aktif/tahun", "Siap uji manual"],
    ["10", "Mapel", "Kode duplikat", "Validasi gagal", "Siap uji manual"],
    ["11", "Jadwal", "Bentrok kelas", "Ditolak", "Lulus otomatis"],
    ["12", "Jadwal", "Bentrok guru", "Ditolak", "Lulus otomatis"],
    ["13", "Nilai", "Input 0 dan 100", "Tersimpan", "Lulus otomatis"],
    ["14", "Nilai", "Nilai kosong", "Tersimpan NULL", "Lulus otomatis"],
    ["15", "Nilai", "Nilai -1/101", "Ditolak", "Lulus otomatis"],
    ["16", "Nilai", "Minggu 5", "Ditolak", "Lulus otomatis"],
    ["17", "Nilai", "Siswa luar kelas", "Ditolak", "Lulus otomatis"],
    ["18", "Siswa", "Melihat nilai sendiri", "Hanya data user", "Lulus otomatis"],
    ["19", "Laporan", "Filter jadwal", "Preview sesuai filter", "Siap uji manual"],
    ["20", "PDF", "Unduh jadwal/nilai", "PDF A4", "Lulus otomatis"],
    ["21", "Print", "Cetak laporan", "Navigasi tersembunyi", "Siap uji manual"],
    ["22", "Mobile", "Input nilai 375px", "Scroll dan sticky", "Siap uji manual"],
]
add_table(["No", "Fitur", "Skenario", "Hasil Diharapkan", "Status"], black_box_rows, 7.5)

add_heading("Lampiran 2. Struktur Basis Data", 2)
add_body("ERD lengkap tersedia pada dokumen teknis docs/erd.dbml. Struktur utama mencakup users, pengaturan_sekolah, tahun_ajaran, semester, guru, siswa, kelas, siswa_kelas, mata_pelajaran, jam_pelajaran, pengajaran, jadwal_pelajaran, nilai_tugas, serta tabel role dan permission.")
add_table(["Relasi", "Kardinalitas/Keterangan"], [
    ["TahunAjaran - Semester", "Satu ke banyak"],
    ["TahunAjaran - Kelas", "Satu ke banyak"],
    ["Siswa - SiswaKelas - Kelas", "Histori banyak penempatan"],
    ["Guru - Pengajaran", "Satu ke banyak"],
    ["Pengajaran - JadwalPelajaran", "Satu ke banyak"],
    ["Pengajaran - NilaiTugas", "Satu ke banyak"],
    ["Siswa - NilaiTugas", "Satu ke banyak"],
], 9)

add_heading("Lampiran 3. Akun Demo Pengembangan", 2)
add_table(["Role", "Username", "Keterangan"], [
    ["Admin", "admin", "Akses penuh"],
    ["Guru", "guru1", "Contoh guru"],
    ["Siswa", "siswa", "Contoh siswa"],
    ["Kepala Sekolah", "kepala", "Monitoring"],
], 9)
add_body("Password demo hanya digunakan pada lingkungan pengembangan dan tidak dicantumkan pada naskah publik. Kredensial production wajib diganti.")

add_heading("Lampiran 4. Panduan Instalasi Ringkas", 2)
installation = [
    "composer install",
    "cp .env.example .env",
    "php artisan key:generate",
    "Atur DB_HOST, DB_DATABASE, DB_USERNAME, dan DB_PASSWORD untuk MySQL",
    "php artisan migrate --seed",
    "php artisan storage:link",
    "npm install",
    "npm run build",
    "php artisan optimize pada production",
]
add_bullets(installation, numbered=True)

add_heading("Lampiran 5. Dokumen Pendukung", 2)
add_bullets([
    "Buku panduan penggunaan: docs/buku-panduan-penggunaan.pdf",
    "Presentasi aplikasi: docs/presentasi-aplikasi-sisd232.pptx",
    "Black box testing: docs/black-box-testing.pdf",
    "Dokumentasi deployment: docs/deployment.md",
    "Dokumentasi role dan permission: docs/roles-permissions.md",
    "Dokumentasi basis data: docs/database.md dan docs/erd.dbml",
])

set_update_fields(doc)

# Metadata and compatibility settings
doc.core_properties.title = TITLE.title()
doc.core_properties.subject = "Skripsi Program Studi S1 Teknik Informatika"
doc.core_properties.author = AUTHOR.title()
doc.core_properties.keywords = "sistem informasi, jadwal pelajaran, nilai siswa, Laravel, Livewire"
doc.core_properties.comments = "Draft akademik lengkap; placeholder institusi dan bukti lapangan wajib dikonfirmasi."

doc.save(OUTPUT)
print(f"Skripsi dibuat: {OUTPUT}")
print(f"Paragraf: {len(doc.paragraphs)}")
print(f"Tabel: {len(doc.tables)}")
