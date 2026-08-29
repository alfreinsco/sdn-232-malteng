from __future__ import annotations

import html
import math
import shutil
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "dokumen-paket-complete"
ASSETS = OUTPUT / "assets"
DATE = "29 Agustus 2026"
SCHOOL = "SD Negeri 232 Maluku Tengah"
SYSTEM = "SISDAR — Sistem Informasi Jadwal Pelajaran dan Nilai Siswa"

NAVY = "12304A"
BLUE = "176B87"
TEAL = "2B9AA0"
PALE = "EAF5F7"
GRAY = "F2F5F7"
LINE = "D9E2E8"
TEXT = "243746"
MUTED = "607585"
WHITE = "FFFFFF"
RED = "B42318"
GREEN = "17835B"
AMBER = "B56A00"

FONT = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


@dataclass
class Spec:
    filename: str
    title: str
    subtitle: str
    code: str
    blocks: list[dict[str, Any]]
    landscape: bool = False


def font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT, size)


def wrapped(draw: ImageDraw.ImageDraw, value: str, width: int, selected_font) -> list[str]:
    words = value.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textlength(candidate, font=selected_font) <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def arrow(draw, start, end, fill="#607585", width=4):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 15
    left = (end[0] - length * math.cos(angle - math.pi / 6), end[1] - length * math.sin(angle - math.pi / 6))
    right = (end[0] - length * math.cos(angle + math.pi / 6), end[1] - length * math.sin(angle + math.pi / 6))
    draw.polygon([end, left, right], fill=fill)


def entity_box(draw, xy, title: str, fields: list[str], accent="#176B87"):
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=14, fill="white", outline="#9FB3C1", width=3)
    draw.rounded_rectangle((x0, y0, x1, y0 + 44), radius=14, fill=accent)
    draw.rectangle((x0, y0 + 28, x1, y0 + 44), fill=accent)
    draw.text((x0 + 14, y0 + 10), title, font=font(20, True), fill="white")
    y = y0 + 54
    for field in fields:
        color = "#12304A" if field.startswith(("PK", "FK", "UK")) else "#3D5363"
        draw.text((x0 + 14, y), field, font=font(15, field.startswith("PK")), fill=color)
        y += 24


def generate_erd() -> Path:
    image = Image.new("RGB", (2600, 1700), "#F4F8FA")
    draw = ImageDraw.Draw(image)
    draw.text((80, 45), "ENTITY RELATIONSHIP DIAGRAM — SISDAR", font=font(36, True), fill="#12304A")
    draw.text((80, 95), "Relasi inti aplikasi akademik dan kontrol akses", font=font(20), fill="#607585")
    boxes = {
        "users": (70, 180, 560, 430),
        "tahun_ajaran": (720, 180, 1190, 390),
        "mata_pelajaran": (1360, 180, 1830, 380),
        "jam_pelajaran": (2000, 180, 2480, 410),
        "guru": (70, 570, 560, 870),
        "semester": (720, 550, 1190, 780),
        "kelas": (1360, 540, 1830, 810),
        "pengajaran": (2000, 540, 2480, 850),
        "siswa": (70, 1050, 560, 1330),
        "siswa_kelas": (720, 1030, 1190, 1280),
        "jadwal_pelajaran": (1360, 1020, 1830, 1280),
        "nilai_tugas": (2000, 1020, 2480, 1380),
        "aktivitas": (70, 1460, 560, 1650),
        "pengaturan_sekolah": (720, 1430, 1190, 1650),
        "rbac": (1360, 1430, 1830, 1650),
    }
    relations = [
        ("users", "guru", "1", "0..1"), ("users", "siswa", "1", "0..1"),
        ("users", "aktivitas", "1", "N"), ("users", "pengaturan_sekolah", "1", "0..1"),
        ("users", "nilai_tugas", "1", "N"), ("users", "rbac", "N", "N"),
        ("tahun_ajaran", "semester", "1", "N"), ("tahun_ajaran", "kelas", "1", "N"),
        ("guru", "kelas", "1", "N"), ("guru", "pengajaran", "1", "N"),
        ("semester", "pengajaran", "1", "N"), ("kelas", "pengajaran", "1", "N"),
        ("mata_pelajaran", "pengajaran", "1", "N"), ("kelas", "siswa_kelas", "1", "N"),
        ("siswa", "siswa_kelas", "1", "N"), ("pengajaran", "jadwal_pelajaran", "1", "N"),
        ("jam_pelajaran", "jadwal_pelajaran", "1", "N"), ("pengajaran", "nilai_tugas", "1", "N"),
        ("siswa", "nilai_tugas", "1", "N"),
    ]
    for source, target, left, right in relations:
        a = boxes[source]
        b = boxes[target]
        ac = ((a[0] + a[2]) // 2, (a[1] + a[3]) // 2)
        bc = ((b[0] + b[2]) // 2, (b[1] + b[3]) // 2)
        if abs(ac[0] - bc[0]) > abs(ac[1] - bc[1]):
            start = (a[2] if bc[0] > ac[0] else a[0], ac[1])
            end = (b[0] if bc[0] > ac[0] else b[2], bc[1])
        else:
            start = (ac[0], a[3] if bc[1] > ac[1] else a[1])
            end = (bc[0], b[1] if bc[1] > ac[1] else b[3])
        draw.line([start, end], fill="#B8C7D1", width=3)
        draw.text((start[0] + 4, start[1] + 4), left, font=font(14, True), fill="#607585")
        draw.text((end[0] - 28, end[1] - 22), right, font=font(14, True), fill="#607585")
    entities = {
        "users": ["PK id", "UK username", "UK email", "name", "password", "status", "last_login_at"],
        "tahun_ajaran": ["PK id", "UK nama", "tanggal_mulai", "tanggal_selesai", "status"],
        "mata_pelajaran": ["PK id", "UK kode", "UK nama", "status"],
        "jam_pelajaran": ["PK id", "nama", "jam_mulai", "jam_selesai", "UK urutan", "jenis", "status"],
        "guru": ["PK id", "FK user_id", "UK nip", "UK nuptk", "nama_lengkap", "jenis_kelamin", "status"],
        "semester": ["PK id", "FK tahun_ajaran_id", "nama", "tanggal_mulai", "tanggal_selesai", "status"],
        "kelas": ["PK id", "FK tahun_ajaran_id", "FK wali_kelas_id", "nama", "tingkat", "status"],
        "pengajaran": ["PK id", "FK semester_id", "FK kelas_id", "FK mata_pelajaran_id", "FK guru_id", "status"],
        "siswa": ["PK id", "FK user_id", "UK nis", "UK nisn", "nama_lengkap", "jenis_kelamin", "status"],
        "siswa_kelas": ["PK id", "FK siswa_id", "FK kelas_id", "status", "UK siswa + kelas"],
        "jadwal_pelajaran": ["PK id", "FK pengajaran_id", "FK jam_pelajaran_id", "hari", "UK pengajaran + hari + jam"],
        "nilai_tugas": ["PK id", "FK pengajaran_id", "FK siswa_id", "FK dibuat_oleh", "bulan (1–12)", "minggu (1–4)", "nilai (0–100 / NULL)", "catatan", "UK pengajaran+siswa+bulan+minggu"],
        "aktivitas": ["PK id", "FK user_id", "subject morph", "type", "description", "properties", "created_at"],
        "pengaturan_sekolah": ["PK id", "FK kepala_sekolah_user_id", "nama_sekolah", "npsn", "alamat", "logo"],
        "rbac": ["roles", "permissions", "model_has_roles", "role_has_permissions", "model_has_permissions"],
    }
    for name, xy in boxes.items():
        entity_box(draw, xy, name, entities[name], "#2B9AA0" if name in {"pengajaran", "jadwal_pelajaran", "nilai_tugas"} else "#176B87")
    path = ASSETS / "erd-sisdar.png"
    image.save(path, quality=95)
    return path


def actor(draw, x, y, label):
    draw.ellipse((x - 22, y - 75, x + 22, y - 31), outline="#12304A", width=4)
    draw.line((x, y - 31, x, y + 35), fill="#12304A", width=4)
    draw.line((x - 40, y - 5, x + 40, y - 5), fill="#12304A", width=4)
    draw.line((x, y + 35, x - 35, y + 85), fill="#12304A", width=4)
    draw.line((x, y + 35, x + 35, y + 85), fill="#12304A", width=4)
    w = draw.textlength(label, font=font(22, True))
    draw.text((x - w / 2, y + 100), label, font=font(22, True), fill="#12304A")


def use_case(draw, xy, label, accent="#176B87"):
    draw.ellipse(xy, fill="white", outline=accent, width=4)
    x0, y0, x1, y1 = xy
    lines = wrapped(draw, label, x1 - x0 - 30, font(19, True))
    total = len(lines) * 25
    y = (y0 + y1 - total) / 2
    for line in lines:
        w = draw.textlength(line, font=font(19, True))
        draw.text(((x0 + x1 - w) / 2, y), line, font=font(19, True), fill="#12304A")
        y += 25


def generate_use_case() -> Path:
    image = Image.new("RGB", (2200, 1400), "#F4F8FA")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "USE CASE DIAGRAM — SISDAR", font=font(36, True), fill="#12304A")
    draw.rounded_rectangle((430, 130, 1770, 1320), radius=24, outline="#9FB3C1", width=4, fill="#FDFEFE")
    draw.text((470, 155), "Batas Sistem SISDAR", font=font(23, True), fill="#607585")
    actor(draw, 210, 360, "Admin")
    actor(draw, 210, 1020, "Guru")
    actor(draw, 1990, 360, "Siswa")
    actor(draw, 1990, 1020, "Kepala Sekolah")
    cases = {
        "auth": (820, 210, 1380, 320, "Login, Logout & Profil"),
        "master": (520, 390, 1040, 510, "Kelola Periode, Master Data & Pengguna"),
        "academic": (1160, 390, 1660, 510, "Kelola Penempatan, Pengajaran & Jadwal"),
        "grade": (520, 620, 1040, 740, "Input dan Monitoring Nilai Mingguan"),
        "view": (1160, 620, 1660, 740, "Lihat Jadwal dan Nilai Pribadi"),
        "report": (520, 850, 1040, 970, "Laporan, Print dan PDF"),
        "monitor": (1160, 850, 1660, 970, "Monitoring Akademik"),
        "activity": (820, 1080, 1380, 1190, "Lihat Aktivitas Pengguna"),
    }
    for _, (x0, y0, x1, y1, label) in cases.items():
        use_case(draw, (x0, y0, x1, y1), label, "#2B9AA0" if label.startswith(("Input", "Laporan", "Monitoring")) else "#176B87")
    connections = [
        ((250, 330), (820, 265)), ((250, 350), (520, 450)), ((250, 380), (1160, 450)), ((250, 410), (520, 680)),
        ((250, 440), (520, 910)), ((250, 470), (820, 1135)),
        ((250, 990), (820, 265)), ((250, 1020), (520, 680)), ((250, 1050), (520, 910)), ((250, 1080), (820, 1135)),
        ((1950, 330), (1380, 265)), ((1950, 360), (1660, 680)), ((1950, 390), (1040, 910)), ((1950, 420), (1380, 1135)),
        ((1950, 990), (1380, 265)), ((1950, 1020), (1660, 910)), ((1950, 1050), (1040, 910)), ((1950, 1080), (1380, 1135)),
    ]
    for start, end in connections:
        draw.line([start, end], fill="#A9BAC6", width=3)
    path = ASSETS / "use-case-sisdar.png"
    image.save(path, quality=95)
    return path


def generate_flow(name: str, title: str, steps: list[tuple[str, str]], decisions: set[int] | None = None) -> Path:
    decisions = decisions or set()
    width = 2200
    height = 360 + max(0, len(steps) - 6) * 80
    image = Image.new("RGB", (width, height), "#F4F8FA")
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), title, font=font(31, True), fill="#12304A")
    margin = 70
    usable = width - margin * 2
    gap = 28
    box_width = (usable - gap * (len(steps) - 1)) / len(steps)
    y0, y1 = 145, 285
    for index, (label, detail) in enumerate(steps):
        x0 = margin + index * (box_width + gap)
        x1 = x0 + box_width
        if index in decisions:
            center = ((x0 + x1) / 2, (y0 + y1) / 2)
            polygon = [(center[0], y0), (x1, center[1]), (center[0], y1), (x0, center[1])]
            draw.polygon(polygon, fill="#FFF4E5", outline="#B56A00")
        else:
            draw.rounded_rectangle((x0, y0, x1, y1), radius=18, fill="white", outline="#176B87", width=4)
        label_lines = wrapped(draw, label, int(box_width - 24), font(19, True))
        detail_lines = wrapped(draw, detail, int(box_width - 24), font(15))
        y = y0 + 18
        for line in label_lines:
            w = draw.textlength(line, font=font(19, True))
            draw.text(((x0 + x1 - w) / 2, y), line, font=font(19, True), fill="#12304A")
            y += 24
        y += 5
        for line in detail_lines[:3]:
            w = draw.textlength(line, font=font(15))
            draw.text(((x0 + x1 - w) / 2, y), line, font=font(15), fill="#607585")
            y += 20
        if index < len(steps) - 1:
            arrow(draw, (x1 + 5, (y0 + y1) / 2), (x1 + gap - 5, (y0 + y1) / 2), width=4)
            if index in decisions:
                draw.text((x1 + 7, (y0 + y1) / 2 - 28), "Ya", font=font(14, True), fill="#17835B")
        if index in decisions:
            branch_x = (x0 + x1) / 2
            draw.line((branch_x, y1, branch_x, y1 + 32), fill="#B42318", width=4)
            draw.rounded_rectangle((branch_x - 105, y1 + 32, branch_x + 105, y1 + 78), radius=12, fill="#FEECEB", outline="#B42318", width=3)
            branch_label = "Tidak: tolak & kembali"
            label_width = draw.textlength(branch_label, font=font(14, True))
            draw.text((branch_x - label_width / 2, y1 + 47), branch_label, font=font(14, True), fill="#B42318")
    path = ASSETS / f"flow-{name}.png"
    image.save(path, quality=95)
    return path


def generate_architecture() -> Path:
    image = Image.new("RGB", (2000, 1050), "#F4F8FA")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "ARSITEKTUR APLIKASI SISDAR", font=font(36, True), fill="#12304A")
    layers = [
        (160, 180, 1840, 330, "PRESENTATION LAYER", "Browser Desktop/Mobile • Blade • Tailwind CSS • JavaScript", "#176B87"),
        (160, 400, 1840, 550, "APPLICATION LAYER", "Laravel Routes • Middleware Auth/Active/Permission • Livewire Components • Controllers", "#2B9AA0"),
        (160, 620, 1840, 770, "DOMAIN & SERVICE LAYER", "Periode Aktif • Validasi Jadwal • Nilai Massal • Penempatan Siswa • Laporan • Aktivitas", "#176B87"),
        (160, 840, 1840, 990, "DATA LAYER", "Eloquent Models • MySQL/MariaDB • Migration • Seeder • Foreign Key • Transaction", "#2B9AA0"),
    ]
    for x0, y0, x1, y1, title, detail, color in layers:
        draw.rounded_rectangle((x0, y0, x1, y1), radius=25, fill="white", outline=color, width=5)
        draw.rounded_rectangle((x0, y0, x0 + 390, y1), radius=25, fill=color)
        draw.rectangle((x0 + 365, y0, x0 + 410, y1), fill=color)
        draw.text((x0 + 35, y0 + 55), title, font=font(25, True), fill="white")
        lines = wrapped(draw, detail, x1 - x0 - 480, font(22))
        y = y0 + 48
        for line in lines:
            draw.text((x0 + 450, y), line, font=font(22), fill="#243746")
            y += 30
    for y in [345, 565, 785]:
        arrow(draw, (1000, y), (1000, y + 40), fill="#607585", width=5)
    path = ASSETS / "arsitektur-sisdar.png"
    image.save(path, quality=95)
    return path


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def set_cell(cell, value: str, header=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    run.font.name = "Aptos"
    run.font.size = Pt(8.1 if not header else 8.4)
    run.font.color.rgb = RGBColor.from_string(WHITE if header else TEXT)
    run.bold = header
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_doc(document: Document, title: str, landscape=False):
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color in [("Title", 25, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 12.5, BLUE), ("Heading 3", 10.8, TEAL)]:
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    for section in document.sections:
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)
        header = section.header.paragraphs[0]
        header.text = title
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)


def doc_cover(document: Document, spec: Spec):
    logo = ROOT / "public" / "logo-malteng.png"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo.exists():
        p.add_run().add_picture(str(logo), height=Inches(1.0))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DOKUMEN PAKET COMPLETE")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(spec.title)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(spec.subtitle)
    r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, row in enumerate([["Sistem", SYSTEM], ["Instansi", SCHOOL], ["Kode Dokumen", spec.code], ["Tanggal", DATE]]):
        shade(table.cell(index, 0), NAVY)
        shade(table.cell(index, 1), PALE)
        set_cell(table.cell(index, 0), row[0], True)
        set_cell(table.cell(index, 1), row[1])
    document.add_page_break()


def render_docx(spec: Spec):
    document = Document()
    configure_doc(document, f"{spec.code} • {spec.title}", spec.landscape)
    doc_cover(document, spec)
    for block in spec.blocks:
        kind = block["type"]
        if kind == "h1":
            document.add_heading(block["text"], level=1)
        elif kind == "h2":
            document.add_heading(block["text"], level=2)
        elif kind == "p":
            document.add_paragraph(block["text"])
        elif kind == "callout":
            result = document.add_table(rows=1, cols=1)
            shade(result.cell(0, 0), block.get("color", PALE))
            set_cell(result.cell(0, 0), f'{block["title"]}\n{block["text"]}')
            document.add_paragraph()
        elif kind in {"bullets", "numbered"}:
            style = "List Bullet" if kind == "bullets" else "List Number"
            for value in block["items"]:
                document.add_paragraph(value, style=style)
        elif kind == "table":
            result = document.add_table(rows=1, cols=len(block["headers"]))
            result.alignment = WD_TABLE_ALIGNMENT.CENTER
            for index, value in enumerate(block["headers"]):
                shade(result.cell(0, index), NAVY)
                set_cell(result.cell(0, index), value, True)
            for row_index, row in enumerate(block["rows"]):
                cells = result.add_row().cells
                for index, value in enumerate(row):
                    shade(cells[index], WHITE if row_index % 2 == 0 else GRAY)
                    set_cell(cells[index], str(value))
            document.add_paragraph()
        elif kind == "image":
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_path = Path(block["path"])
            width = block.get("width", 6.7 if not spec.landscape else 9.7)
            p.add_run().add_picture(str(image_path), width=Inches(width))
            if block.get("caption"):
                c = document.add_paragraph(block["caption"])
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in c.runs:
                    run.italic = True
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor.from_string(MUTED)
        elif kind == "pagebreak":
            document.add_page_break()
    document.save(OUTPUT / f"{spec.filename}.docx")


def html_table(headers, rows):
    head = "".join(f"<th>{html.escape(str(value))}</th>" for value in headers)
    body = "".join("<tr>" + "".join(f"<td>{html.escape(str(value))}</td>" for value in row) + "</tr>" for row in rows)
    return f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>"


def render_html(spec: Spec):
    body = []
    for block in spec.blocks:
        kind = block["type"]
        if kind in {"h1", "h2"}:
            body.append(f"<{kind}>{html.escape(block['text'])}</{kind}>")
        elif kind == "p":
            body.append(f"<p>{html.escape(block['text'])}</p>")
        elif kind == "callout":
            body.append(f"<div class='callout'><strong>{html.escape(block['title'])}</strong>{html.escape(block['text'])}</div>")
        elif kind in {"bullets", "numbered"}:
            tag = "ul" if kind == "bullets" else "ol"
            body.append(f"<{tag}>" + "".join(f"<li>{html.escape(value)}</li>" for value in block["items"]) + f"</{tag}>")
        elif kind == "table":
            body.append(html_table(block["headers"], block["rows"]))
        elif kind == "image":
            uri = Path(block["path"]).resolve().as_uri()
            caption = f"<figcaption>{html.escape(block.get('caption', ''))}</figcaption>" if block.get("caption") else ""
            body.append(f"<figure><img src='{uri}'>{caption}</figure>")
        elif kind == "pagebreak":
            body.append("<div class='pagebreak'></div>")
    logo = (ROOT / "public" / "logo-malteng.png").resolve().as_uri()
    page_size = "A4 landscape" if spec.landscape else "A4"
    content = f"""<!doctype html><html lang="id"><head><meta charset="utf-8"><title>{html.escape(spec.title)}</title><style>
@page {{ size:{page_size}; margin:15mm 15mm 17mm; }}
:root{{--navy:#12304a;--blue:#176b87;--teal:#2b9aa0;--ink:#243746;--muted:#607585;--line:#d9e2e8;--pale:#eaf5f7;--gray:#f2f5f7;}}
*{{box-sizing:border-box}} body{{margin:0;color:var(--ink);font:10pt/1.48 Arial,sans-serif}} .cover{{min-height:{'174mm' if spec.landscape else '258mm'};display:flex;flex-direction:column;justify-content:center;text-align:center;page-break-after:always}}
.cover img{{width:24mm;height:29mm;object-fit:contain;margin:0 auto 8mm}} .eyebrow{{color:var(--teal);font-size:9pt;font-weight:700;letter-spacing:1.7px}} .cover h1{{margin:3mm 0;color:var(--navy);font-size:28pt;border:0}}
.subtitle{{color:var(--blue);font-size:13pt;max-width:160mm;margin:0 auto 9mm}} .meta{{width:160mm;margin:0 auto;border-collapse:separate;border-spacing:0 1.5mm;text-align:left}} .meta th,.meta td{{padding:3mm 4mm}} .meta th{{width:42mm;color:white;background:var(--navy)}} .meta td{{background:var(--pale);font-weight:600}}
h1{{margin:8mm 0 3mm;padding-bottom:2mm;color:var(--navy);border-bottom:1.1mm solid var(--teal);font-size:18pt;break-after:avoid}} h1:first-child{{margin-top:0}} h2{{margin:6mm 0 2mm;color:var(--blue);font-size:13pt;break-after:avoid}} p{{margin:0 0 3mm}} ul,ol{{margin:1.5mm 0 4mm;padding-left:6mm}} li{{margin-bottom:1.3mm}}
.callout{{margin:3mm 0 5mm;padding:4mm 5mm;border-left:1.5mm solid var(--teal);background:var(--pale);border-radius:1.5mm}} .callout strong{{display:block;color:var(--navy);font-size:11pt;margin-bottom:1mm}}
table{{width:100%;margin:3mm 0 5mm;border-collapse:collapse;font-size:8pt}} thead{{display:table-header-group}} tr{{break-inside:avoid}} th{{padding:2.1mm;color:white;background:var(--navy);text-align:left}} td{{padding:2mm;border:.2mm solid var(--line);vertical-align:top}} tbody tr:nth-child(even) td{{background:var(--gray)}}
figure{{margin:4mm 0 6mm;text-align:center;break-inside:avoid}} figure img{{max-width:100%;max-height:{'120mm' if spec.landscape else '225mm'};object-fit:contain;border:.2mm solid var(--line)}} figcaption{{margin-top:1.5mm;color:var(--muted);font-size:8.5pt;font-style:italic}} .pagebreak{{page-break-before:always}}
</style></head><body><section class="cover"><img src="{logo}"><div class="eyebrow">DOKUMEN PAKET COMPLETE</div><h1>{html.escape(spec.title)}</h1><p class="subtitle">{html.escape(spec.subtitle)}</p><table class="meta"><tr><th>Sistem</th><td>{html.escape(SYSTEM)}</td></tr><tr><th>Instansi</th><td>{html.escape(SCHOOL)}</td></tr><tr><th>Kode Dokumen</th><td>{html.escape(spec.code)}</td></tr><tr><th>Tanggal</th><td>{DATE}</td></tr></table></section><main>{''.join(body)}</main></body></html>"""
    (OUTPUT / f"{spec.filename}.html").write_text(content, encoding="utf-8")


def h1(text): return {"type": "h1", "text": text}
def h2(text): return {"type": "h2", "text": text}
def p(text): return {"type": "p", "text": text}
def bullets(items): return {"type": "bullets", "items": items}
def numbered(items): return {"type": "numbered", "items": items}
def table(headers, rows): return {"type": "table", "headers": headers, "rows": rows}
def image(path, caption, width=None):
    value = {"type": "image", "path": str(path), "caption": caption}
    if width: value["width"] = width
    return value
def callout(title, text): return {"type": "callout", "title": title, "text": text}


def build_specs(diagrams: dict[str, Path]) -> list[Spec]:
    entity_rows = [
        ["users", "Akun login dan identitas pengguna", "username, email", "guru, siswa, aktivitas, pengaturan, nilai, RBAC"],
        ["pengaturan_sekolah", "Identitas sekolah dan Kepala Sekolah", "—", "users"],
        ["tahun_ajaran", "Periode tahun akademik", "nama", "semester, kelas"], ["semester", "Semester per tahun ajaran", "tahun_ajaran + nama", "pengajaran"],
        ["guru", "Profil guru dan relasi akun", "user_id, nip, nuptk", "kelas, pengajaran"], ["siswa", "Profil siswa dan relasi akun", "user_id, nis, nisn", "siswa_kelas, nilai_tugas"],
        ["kelas", "Rombongan belajar per tahun ajaran", "tahun_ajaran + nama", "siswa_kelas, pengajaran"], ["siswa_kelas", "Histori penempatan siswa", "siswa + kelas", "siswa, kelas"],
        ["mata_pelajaran", "Master mata pelajaran", "kode, nama", "pengajaran"], ["jam_pelajaran", "Slot waktu dan urutan", "urutan", "jadwal_pelajaran"],
        ["pengajaran", "Guru mengajar mapel di kelas/semester", "semester+kelas+mapel+guru", "jadwal, nilai"],
        ["jadwal_pelajaran", "Hari dan jam untuk pengajaran", "pengajaran+hari+jam", "pengajaran, jam"],
        ["nilai_tugas", "Nilai per siswa, bulan, dan minggu", "pengajaran+siswa+bulan+minggu", "pengajaran, siswa, users"],
        ["aktivitas", "Audit aktivitas penting pengguna", "—", "users dan subject polymorphic"], ["RBAC", "Role dan permission", "kombinasi role/permission/model", "users"],
    ]
    erd = Spec("01-ERD-SISDAR", "ENTITY RELATIONSHIP DIAGRAM", "Struktur basis data, relasi, dan aturan integritas SISDAR", "PC-01/ERD", [
        h1("1. Tujuan Dokumen"), p("Dokumen ini menjelaskan struktur basis data SISDAR, hubungan antarentitas, kunci utama/asing, serta aturan integritas yang menjaga histori akademik dan keamanan akses data."),
        h1("2. Diagram ERD"), image(diagrams["erd"], "Gambar 1. ERD inti SISDAR. Notasi 1 menunjukkan satu record dan N menunjukkan banyak record.", 8.2),
        h1("3. Kamus Entitas"), table(["Entitas", "Fungsi", "Kunci Unik Penting", "Relasi Utama"], entity_rows),
        h1("4. Aturan Relasi dan Integritas"), bullets(["Foreign key akademik menggunakan restrict-on-delete agar histori tidak terhapus secara tidak sengaja.", "Relasi akun guru/siswa dan Kepala Sekolah bersifat opsional serta menggunakan null-on-delete.", "Satu tahun ajaran dan satu semester aktif pada satu waktu melalui service aktivasi transactional.", "Penempatan siswa disimpan pada siswa_kelas sehingga histori kelas per tahun ajaran tetap tersedia.", "Jadwal divalidasi terhadap bentrok kelas, bentrok guru, serta duplikasi pengajaran-hari-jam.", "Nilai unik untuk kombinasi pengajaran, siswa, bulan, dan minggu; nilai NULL berarti belum dinilai sedangkan 0 adalah nilai sah.", "Urutan jam pelajaran unik dan dikelola melalui reorder transactional.", "Role dan permission disimpan melalui tabel RBAC Spatie Laravel Permission."]),
        h1("5. Relasi Kritis"), table(["Induk", "Anak", "Kardinalitas", "Makna"], [["tahun_ajaran", "semester", "1 : N", "Satu tahun ajaran memiliki beberapa semester."], ["tahun_ajaran", "kelas", "1 : N", "Kelas terikat pada tahun ajaran tertentu."], ["kelas", "siswa_kelas", "1 : N", "Satu kelas memiliki banyak histori anggota."], ["siswa", "siswa_kelas", "1 : N", "Siswa dapat memiliki histori kelas lintas periode."], ["semester/kelas/mapel/guru", "pengajaran", "1 : N", "Pengajaran menjadi pusat transaksi akademik."], ["pengajaran", "jadwal_pelajaran", "1 : N", "Satu pengajaran dapat muncul pada beberapa slot jadwal."], ["pengajaran+siswa", "nilai_tugas", "1 : N", "Nilai disimpan per bulan dan minggu."], ["users", "aktivitas", "1 : N", "Aktivitas penting dicatat untuk audit."]]),
    ], landscape=True)

    usecase_rows = [["UC-01", "Login dan logout", "Semua role", "Akun aktif dan kredensial valid", "Session dibuat/dihapus"], ["UC-02", "Kelola periode dan master data", "Admin", "Login sebagai Admin", "Data master tersimpan/berstatus"], ["UC-03", "Penempatan siswa", "Admin", "Siswa dan kelas tersedia", "Keanggotaan aktif dan histori terjaga"], ["UC-04", "Kelola pengajaran dan jadwal", "Admin", "Periode/master lengkap", "Jadwal tersimpan tanpa bentrok"], ["UC-05", "Input nilai mingguan", "Guru", "Memiliki pengajaran", "Nilai M1–M4 tersimpan massal"], ["UC-06", "Lihat jadwal/nilai pribadi", "Siswa", "Terhubung ke profil siswa", "Hanya data sendiri tampil"], ["UC-07", "Monitoring akademik", "Admin/Kepala Sekolah", "Login dan permission", "Ringkasan jadwal/nilai tampil"], ["UC-08", "Print dan ekspor PDF", "Semua role sesuai data", "Filter dan permission valid", "Laporan sesuai scope role"], ["UC-09", "Profil dan aktivitas", "Semua role", "Login", "Profil terbarui dan aktivitas terlihat sesuai scope"]]
    useflow = Spec("02-Use-Case-dan-Flow-Sistem-SISDAR", "USE CASE DAN FLOW SISTEM", "Aktor, interaksi, dan alur proses utama SISDAR", "PC-02/UCF", [
        h1("1. Aktor Sistem"), table(["Aktor", "Tanggung Jawab"], [["Admin", "Mengelola periode, master data, pengguna, penempatan siswa, pengajaran, jadwal, monitoring, laporan, dan pengaturan sekolah."], ["Guru", "Melihat jadwal mengajar, menginput nilai pengajaran sendiri, membuka laporan, profil, dan aktivitas sendiri."], ["Siswa", "Melihat jadwal kelas, nilai sendiri, laporan pribadi, profil, dan aktivitas sendiri."], ["Kepala Sekolah", "Melakukan monitoring read-only terhadap dashboard, guru, siswa, kelas, jadwal, nilai, laporan, dan aktivitas."]]),
        h1("2. Use Case Diagram"), image(diagrams["usecase"], "Gambar 1. Use case utama dan hubungan empat aktor SISDAR."),
        h1("3. Daftar Use Case"), table(["ID", "Use Case", "Aktor", "Prasyarat", "Hasil"], usecase_rows),
        h1("4. Flow Autentikasi"), image(diagrams["login"], "Gambar 2. Alur login, pemeriksaan status akun, role, dan dashboard."),
        h1("5. Flow Setup Akademik"), image(diagrams["setup"], "Gambar 3. Urutan setup periode, master data, penempatan, pengajaran, dan jadwal."),
        h1("6. Flow Penyusunan Jadwal"), image(diagrams["schedule"], "Gambar 4. Validasi bentrok sebelum jadwal disimpan."),
        h1("7. Flow Input Nilai"), image(diagrams["grade"], "Gambar 5. Input nilai massal dan validasi hak guru serta anggota kelas."),
        h1("8. Flow Laporan"), image(diagrams["report"], "Gambar 6. Preview, print, dan ekspor PDF dengan pembatasan sesuai role."),
        h1("9. Aturan Keputusan Penting"), bullets(["Akun nonaktif tidak dapat login.", "Semester hanya dapat aktif jika tahun ajarannya aktif.", "Jadwal ditolak apabila kelas atau guru bentrok pada semester, hari, dan jam yang sama.", "Guru hanya dapat memproses nilai pada pengajaran miliknya.", "Siswa ditentukan dari relasi akun login, bukan ID yang dikirim browser.", "Laporan guru dan siswa otomatis dibatasi oleh identitas akun."]),
    ])

    blackbox_rows = [line.split("|")[1:-1] for line in """|1|Login|Kredensial benar|Masuk dashboard sesuai role|Lulus otomatis|
|2|Login|Password salah|Ditolak dengan pesan validasi|Lulus otomatis|
|3|Login|Akun nonaktif|Login ditolak|Lulus otomatis|
|4|Registrasi|Akses /register|Halaman tidak tersedia|Lulus otomatis|
|5|Authorization|Siswa membuka /pengguna|HTTP 403|Lulus otomatis|
|6|Authorization|Guru membuka tahun ajaran|HTTP 403|Lulus otomatis|
|7|Guru|Tambah/ubah data valid|Data tersimpan|Siap uji manual|
|8|Siswa|NIS/NISN duplikat|Validasi gagal|Siap uji manual|
|9|Kelas|Penempatan banyak siswa|Satu kelas aktif per tahun|Siap uji manual|
|10|Mata Pelajaran|Kode duplikat|Validasi gagal|Siap uji manual|
|11|Jadwal|Bentrok kelas|Ditolak|Lulus otomatis|
|12|Jadwal|Bentrok guru|Ditolak|Lulus otomatis|
|13|Nilai|Nilai 0 dan 100|Keduanya tersimpan|Lulus otomatis|
|14|Nilai|Nilai kosong|Disimpan NULL|Lulus otomatis|
|15|Nilai|Nilai -1 atau 101|Ditolak|Lulus otomatis|
|16|Nilai|Minggu ke-5|Ditolak|Lulus otomatis|
|17|Nilai|Siswa di luar kelas|Ditolak|Lulus otomatis|
|18|Siswa|Melihat nilai sendiri|Hanya nilai akun login|Lulus otomatis|
|19|Laporan|Filter jadwal|Preview sesuai filter|Siap uji manual|
|20|PDF|Unduh laporan|File PDF A4 terunduh|Lulus otomatis|
|21|Print|Cetak laporan|Navigasi/filter tersembunyi|Siap uji manual|
|22|Mobile|Input nilai 375px|Tabel dapat digeser|Siap uji manual|
|23|Jam Pelajaran|Tambah tanpa urutan|Urutan otomatis|Lulus otomatis|
|24|Jam Pelajaran|Drag-and-drop urutan|Urutan tersimpan aman|Lulus otomatis|
|25|Status Master|Nonaktifkan record|Status berubah, data tidak terhapus|Lulus otomatis|
|26|Tabel|Select all lintas halaman|Semua hasil filter terpilih|Lulus otomatis|
|27|Tabel|Kecualikan satu baris|Jumlah selection berkurang|Lulus otomatis|
|28|Tabel|Refresh state URL|State aman dipulihkan|Lulus otomatis|
|29|Pagination|Nilai di luar batas|Dibatasi ke rentang aman|Lulus otomatis|
|30|Nilai/Laporan|Periksa aksi tabel|Checkbox/aksi tidak tampil|Lulus otomatis|
|31|Pengajaran|Filter bertingkat|Hasil dan URL sesuai|Lulus otomatis|
|32|Jadwal|Pilihan bertingkat|Pengajaran relevan tampil|Lulus otomatis|
|33|Kelas|Lihat anggota kelas|Roster tampil|Lulus otomatis|
|34|Kelas|Tambah/keluarkan siswa|Keanggotaan berubah, siswa tetap ada|Lulus otomatis|
|35|Penempatan|Filter status kelas|Hasil sesuai status|Lulus otomatis|
|36|Dashboard|Render seluruh role|Metrik sesuai role|Lulus otomatis|
|37|Mobile|Kolom aksi 390px|Data tidak tertutup|Siap uji manual|""".strip().splitlines()]
    blackbox = Spec("03-Black-Box-Testing-SISDAR", "BLACK BOX TESTING", "Skenario uji fungsi, hasil aktual, dan catatan penerimaan", "PC-03/BBT", [
        h1("1. Tujuan dan Metode"), p("Pengujian black box memeriksa perilaku aplikasi dari sisi masukan, proses yang terlihat, keluaran, validasi, dan pembatasan akses tanpa menilai detail implementasi internal."),
        h1("2. Lingkungan dan Hasil Verifikasi"), table(["Pemeriksaan", "Tanggal", "Hasil"], [["Automated test PHPUnit", DATE, "66 test dijalankan; 64 lulus, 2 gagal; 1.709 assertion."], ["Production frontend build", DATE, "Berhasil menggunakan Vite 8.2.1."], ["Pemeriksaan PDF/dokumen", DATE, "Dilakukan saat ekspor dokumen Paket Complete."]]),
        callout("Catatan tindak lanjut", "Dua kegagalan automated test berada pada LandingPageTest: ekspektasi redirect pengguna login dari landing page dan akses landing melalui riwayat/back button. Modul akademik inti, authorization, jadwal, nilai, tabel, dan PDF tetap lulus pada suite terkait. Dua kasus landing harus diselesaikan sebelum status penerimaan final dinyatakan tanpa catatan."),
        h1("3. Skenario Black Box"), table(["No.", "Fitur", "Skenario", "Hasil yang Diharapkan", "Status"], blackbox_rows),
        h1("4. Prosedur Uji Manual"), numbered(["Siapkan database demo dan pastikan tahun ajaran/semester aktif.", "Uji setiap role menggunakan browser desktop dan mobile/responsive mode.", "Catat data awal, tindakan, hasil aktual, bukti screenshot, dan status.", "Untuk kasus perubahan data, kembalikan data demo setelah pengujian.", "Tandai Lulus jika hasil aktual sama dengan hasil yang diharapkan; selain itu buat catatan defect."]),
        h1("5. Klasifikasi Hasil"), table(["Status", "Arti"], [["Lulus otomatis", "Telah dibuktikan oleh automated test terkait."], ["Siap uji manual", "Memerlukan verifikasi visual/interaksi bersama pengguna."], ["Perlu tindak lanjut", "Ditemukan ketidaksesuaian yang harus diperbaiki atau dikonfirmasi."]]),
        h1("6. Lembar Penerimaan Pengujian"), table(["Pihak", "Nama", "Tanggal", "Tanda Tangan"], [["Client/Penerima", "", "", ""], ["Pengembang", "", "", ""]]),
    ])

    documentation = Spec("04-Dokumentasi-Sistem-SISDAR", "DOKUMENTASI SISTEM", "Arsitektur, modul, keamanan, instalasi, deployment, dan pemeliharaan", "PC-04/DKS", [
        h1("1. Ikhtisar Sistem"), p("SISDAR adalah aplikasi web untuk pengelolaan jadwal pelajaran dan nilai siswa pada SD Negeri 232 Maluku Tengah. Sistem mendukung Admin, Guru, Siswa, dan Kepala Sekolah dengan hak akses terpisah."),
        h1("2. Arsitektur"), image(diagrams["architecture"], "Gambar 1. Arsitektur berlapis SISDAR."),
        h1("3. Teknologi"), table(["Komponen", "Implementasi"], [["Backend", "PHP 8.3+, Laravel 13"], ["UI interaktif", "Livewire 4 single-file components, Blade"], ["Styling/build", "Tailwind CSS 4, Vite 8"], ["Database", "MySQL/MariaDB; SQLite in-memory untuk test"], ["Authorization", "Spatie Laravel Permission"], ["PDF", "DomPDF"], ["Testing", "PHPUnit 12"]]),
        h1("4. Modul Aplikasi"), table(["Modul", "Fungsi"], [["Autentikasi", "Login/logout, status akun, session, dan role."], ["Dashboard", "Metrik role-aware, periode aktif, progres nilai, jadwal hari ini."], ["Master Data", "Tahun ajaran, semester, guru, siswa, kelas, mapel, jam, pengguna."], ["Penempatan", "Anggota kelas dan histori penempatan siswa."], ["Pengajaran", "Relasi guru-mapel-kelas-semester."], ["Jadwal", "Penyusunan dan validasi anti-bentrok."], ["Nilai", "Input massal M1–M4 per bulan dan rata-rata."], ["Laporan", "Preview web, print, dan PDF jadwal/nilai."], ["Pengaturan", "Identitas sekolah dan Kepala Sekolah."], ["Aktivitas/Profil", "Audit kegiatan penting dan profil akun."]]),
        h1("5. Hak Akses"), table(["Area", "Admin", "Guru", "Siswa", "Kepala Sekolah"], [["Dashboard", "Penuh", "Pribadi", "Pribadi", "Monitoring"], ["Master data", "Kelola", "Baca terbatas", "—", "Baca"], ["Pengajaran", "Kelola", "Sendiri", "—", "Baca"], ["Jadwal", "Kelola", "Mengajar", "Kelas", "Semua"], ["Nilai", "Monitoring", "Input sendiri", "Pribadi", "Monitoring"], ["Laporan", "Semua", "Pengajaran", "Pribadi", "Semua"]]),
        h1("6. Aturan Bisnis"), bullets(["Satu tahun ajaran dan semester aktif pada satu waktu.", "Guru dan kelas tidak boleh bentrok pada semester, hari, dan jam yang sama.", "Nilai disimpan per pengajaran, siswa, bulan 1–12, dan Minggu 1–4.", "Nilai kosong adalah NULL; nilai 0 tetap sah; rata-rata hanya memakai nilai terisi.", "Guru hanya mengelola nilai pengajarannya; siswa hanya melihat data sendiri.", "Data historis dipertahankan melalui status aktif/nonaktif dan foreign key restriktif.", "Registrasi publik dinonaktifkan."]),
        h1("7. Struktur Proyek"), table(["Direktori", "Isi"], [["app/Models", "Model dan relasi Eloquent."], ["app/Services", "Business logic periode, jadwal, nilai, laporan, aktivitas."], ["resources/views/pages", "Halaman Livewire per modul."], ["resources/views/components", "Komponen navigasi, select, dan data table."], ["routes/web.php", "Route dan middleware akses."], ["database/migrations", "Struktur database."], ["database/seeders", "Role, permission, dan data demo."], ["tests/Feature", "Automated feature tests."], ["docs", "Panduan dan dokumentasi."]]),
        h1("8. Keamanan"), bullets(["Route dilindungi middleware auth, active, role, dan permission.", "Query nilai/laporan dibatasi lagi berdasarkan user login untuk mencegah IDOR.", "Password disimpan dalam hash oleh Laravel.", "Validasi server digunakan untuk field, rentang nilai, kepemilikan pengajaran, dan anggota kelas.", "Transaksi database digunakan pada aktivasi periode, reorder jam, penempatan, dan penyimpanan nilai massal.", "Production wajib menggunakan APP_DEBUG=false, HTTPS, secure cookie, dan credential database berhak minimum."]),
        h1("9. Instalasi"), numbered(["Jalankan composer install.", "Salin .env.example menjadi .env dan isi database.", "Jalankan php artisan key:generate.", "Jalankan php artisan migrate --seed.", "Jalankan php artisan storage:link.", "Jalankan npm install dan npm run build.", "Jalankan php artisan serve untuk development atau arahkan web server ke public/ untuk production."]),
        h1("10. Deployment Production"), bullets(["Gunakan PHP 8.3+, Composer 2, MySQL 8+/MariaDB, Nginx/Apache, dan HTTPS.", "Set APP_ENV=production, APP_DEBUG=false, APP_URL HTTPS, APP_TIMEZONE=Asia/Jayapura, dan SESSION_SECURE_COOKIE=true.", "Berikan write permission hanya pada storage/ dan bootstrap/cache/.", "Jalankan composer install --no-dev, npm run build, php artisan migrate --force, dan php artisan optimize.", "Siapkan backup database dan storage/app/public sebelum pembaruan."]),
        h1("11. Pemeliharaan dan Troubleshooting"), table(["Masalah", "Pemeriksaan"], [["Tidak dapat login", "Periksa username, status akun, password, session, dan log."], ["Data tidak muncul", "Periksa periode aktif, filter, status record, serta relasi user-profil."], ["Jadwal ditolak", "Periksa bentrok guru/kelas dan kombinasi hari-jam."], ["Siswa tidak muncul di nilai", "Periksa penempatan kelas aktif dan pengajaran."], ["PDF gagal", "Periksa permission laporan, data sekolah, logo, dan writable temp/storage."], ["Aset tampilan hilang", "Jalankan npm run build dan periksa public/build/manifest.json."]]),
    ])

    demo = Spec("05-Persiapan-Demonstrasi-SISDAR", "PERSIAPAN DEMONSTRASI", "Runbook presentasi, akun, data, urutan demo, dan rencana pemulihan", "PC-05/DEM", [
        h1("1. Tujuan"), p("Dokumen ini menjadi panduan operasional agar demonstrasi SISDAR berjalan terstruktur, konsisten, dan dapat dipulihkan dengan cepat apabila terjadi kendala."),
        h1("2. Checklist Sebelum Demonstrasi"), bullets(["Laptop dan charger tersedia; browser terbaru telah diuji.", "Aplikasi dapat dibuka melalui URL lokal/hosting yang akan digunakan.", "Database demo sudah terisi dan tahun ajaran 2026/2027 serta semester ganjil aktif.", "Akun Admin, Guru, Siswa, dan Kepala Sekolah telah diuji login.", "Logo dan identitas sekolah tampil pada laporan.", "PDF dapat diunduh dan folder download mudah ditemukan.", "Koneksi cadangan/hotspot tersedia bila demo memakai hosting.", "Presentasi PPTX/PDF dan screenshot tersedia sebagai fallback offline."]),
        h1("3. Akun Demonstrasi"), table(["Role", "Username", "Password Seeder Saat Ini", "Catatan"], [["Admin", "admin", "123", "Kelola seluruh modul."], ["Guru", "guru1", "123", "Gunakan pengajaran yang memiliki siswa dan nilai."], ["Siswa", "siswa", "123", "Tampilkan jadwal dan nilai pribadi."], ["Kepala Sekolah", "kepala", "123", "Monitoring read-only."]]),
        callout("Keamanan akun demo", "Password 123 berasal dari seeder aplikasi saat ini dan hanya layak untuk lingkungan demonstrasi lokal. Ganti seluruh password dengan password kuat sebelum deployment atau sebelum aplikasi dapat diakses publik. Jangan tampilkan password pada slide utama."),
        h1("4. Data yang Harus Tersedia"), table(["Data", "Kondisi Minimum"], [["Tahun ajaran/semester", "Satu tahun ajaran dan semester aktif."], ["Guru", "Minimal satu guru terhubung akun demo."], ["Siswa", "Minimal satu siswa terhubung akun dan ditempatkan di kelas."], ["Kelas", "Memiliki wali kelas dan anggota."], ["Mata pelajaran/jam", "Tersedia untuk membuat pengajaran dan jadwal."], ["Pengajaran", "Guru-mapel-kelas-semester lengkap."], ["Jadwal", "Ada jadwal valid serta data untuk simulasi bentrok."], ["Nilai", "Ada nilai M1–M4 sebagian/terisi untuk menunjukkan rata-rata."], ["Pengaturan sekolah", "Nama sekolah, alamat, Kepala Sekolah, dan logo terisi."]]),
        h1("5. Susunan Demonstrasi 15–20 Menit"), table(["Durasi", "Role/Halaman", "Poin yang Ditunjukkan"], [["1 menit", "Landing & Login", "Identitas sistem dan empat role."], ["3 menit", "Admin Dashboard", "Metrik, periode aktif, progres nilai, akses cepat."], ["4 menit", "Admin Master/Kelas", "Tabel, filter, anggota kelas, jam drag-and-drop."], ["3 menit", "Admin Jadwal", "Pilihan bertingkat dan penolakan bentrok."], ["3 menit", "Guru Nilai", "Input massal M1–M4, nilai 0 vs kosong, rata-rata."], ["2 menit", "Siswa", "Jadwal kelas dan nilai pribadi."], ["2 menit", "Kepala Sekolah", "Monitoring read-only."], ["2 menit", "Laporan", "Filter, print, dan PDF."]]),
        h1("6. Langkah Demonstrasi Utama"), numbered(["Login Admin dan jelaskan dashboard role-aware.", "Buka Kelas, tunjukkan jumlah siswa, lalu halaman Anggota Kelas.", "Buka Jam Pelajaran dan tunjukkan urutan drag-and-drop.", "Buka Pengajaran/Jadwal dan jelaskan pilihan tahun ajaran → semester → kelas.", "Simulasikan data jadwal bentrok tanpa menyimpan perubahan permanen.", "Login Guru, pilih pengajaran dan bulan, lalu jelaskan input massal nilai M1–M4.", "Login Siswa dan buktikan hanya jadwal/nilai pribadi yang terlihat.", "Login Kepala Sekolah dan tunjukkan monitoring tanpa tombol perubahan data.", "Buka laporan, gunakan filter, print preview, lalu unduh PDF."]),
        h1("7. Bukti Tampilan"), image(ROOT / "docs/screenshots/admin-dashboard.png", "Dashboard Admin sebagai pembuka demonstrasi."), image(ROOT / "docs/screenshots/guru-input-nilai-terisi.png", "Input nilai massal Guru."), image(ROOT / "docs/screenshots/kepala-laporan-nilai.png", "Monitoring dan laporan Kepala Sekolah."),
        h1("8. Rencana Pemulihan"), table(["Kendala", "Tindakan Cepat"], [["Internet putus", "Gunakan aplikasi lokal atau presentasi PDF dan screenshot."], ["Login gagal", "Periksa status akun, gunakan akun alternatif, atau reset database demo."], ["Data demo berubah", "Jalankan migrate:fresh --seed pada lingkungan demo yang aman."], ["PDF tidak terbuka", "Gunakan laporan web/print dan file PDF contoh."], ["Browser lambat", "Tutup tab lain, reload, atau gunakan browser cadangan."], ["Pertanyaan di luar scope", "Catat sebagai pengembangan lanjutan, jangan menjanjikan implementasi langsung."]]),
        h1("9. Pertanyaan yang Perlu Disiapkan"), bullets(["Mengapa jadwal tidak dapat bentrok?", "Mengapa nilai kosong berbeda dari nol?", "Bagaimana siswa naik kelas tanpa kehilangan histori?", "Bagaimana hak akses siswa dilindungi?", "Apakah laporan dapat dicetak dan diunduh?", "Bagaimana proses deployment dan backup?", "Fitur apa yang berada di luar ruang lingkup saat ini?"]),
        h1("10. Checklist Setelah Demonstrasi"), bullets(["Catat seluruh pertanyaan dan revisi minor.", "Kembalikan data demo ke kondisi bersih bila ada perubahan.", "Simpan bukti PDF/screenshot hasil demonstrasi.", "Konfirmasi item yang diterima, perlu revisi, atau di luar scope.", "Sepakati jadwal tindak lanjut dan media komunikasi."]),
    ])

    support = Spec("06-Pendampingan-Teknis-SISDAR", "PENDAMPINGAN TEKNIS", "Rencana sesi, SOP operasional, deployment, troubleshooting, dan serah pengetahuan", "PC-06/PTK", [
        h1("1. Tujuan dan Batas Pendampingan"), p("Pendampingan teknis membantu penerima memahami struktur aplikasi, database, alur CRUD, role-permission, jadwal, nilai, laporan, deployment, dan pemeliharaan dasar. Pendampingan tidak otomatis mencakup pengembangan fitur baru di luar ruang lingkup Paket Complete."),
        h1("2. Peserta yang Disarankan"), bullets(["Admin/operator sekolah sebagai pengguna operasional utama.", "Guru perwakilan untuk alur jadwal dan nilai.", "Kepala Sekolah untuk monitoring dan laporan.", "Pengelola hosting/teknis bila aplikasi akan dipasang di server.", "Pemilik projek akhir untuk penjelasan arsitektur dan demonstrasi."]),
        h1("3. Rencana Sesi"), table(["Sesi", "Durasi", "Materi", "Hasil"], [["1. Orientasi Sistem", "60 menit", "Arsitektur, modul, role, dan navigasi", "Peserta memahami gambaran umum."], ["2. Operasional Admin", "90 menit", "Periode, master data, penempatan, pengajaran, jadwal", "Admin mampu setup akademik."], ["3. Guru dan Nilai", "60 menit", "Jadwal mengajar, input M1–M4, laporan", "Guru mampu menginput dan mengecek nilai."], ["4. Siswa/Kepala", "45 menit", "Akses pribadi dan monitoring", "Aktor memahami batas akses."], ["5. Teknis & Deployment", "90 menit", "Instalasi, .env, database, build, HTTPS, backup", "Tim teknis mampu menjalankan aplikasi."], ["6. Evaluasi", "30 menit", "Tanya jawab dan praktik mandiri", "Daftar tindak lanjut tersedia."]]),
        h1("4. Materi Penjelasan Teknis"), table(["Topik", "Pokok Bahasan"], [["Struktur aplikasi", "Laravel MVC, Livewire, routes, models, services, views, dan tests."], ["Database", "ERD, foreign key, unique index, histori, dan transaksi."], ["CRUD", "Validasi, status aktif/nonaktif, bulk action, dan audit aktivitas."], ["Role/permission", "Admin, Guru, Siswa, Kepala Sekolah, middleware, dan query scope."], ["Jadwal", "Pengajaran, slot jam, validasi bentrok kelas/guru."], ["Nilai", "Bulk save, bulan, Minggu 1–4, NULL, 0, rata-rata, authorization."], ["Laporan", "Filter role-aware, print stylesheet, dan PDF."], ["Deployment", "Environment production, web root public, permission folder, HTTPS, optimize."]]),
        h1("5. SOP Operasional Ringkas"), numbered(["Admin memastikan tahun ajaran dan semester aktif.", "Admin memeriksa master data dan status akun.", "Admin menempatkan siswa ke kelas dan membuat pengajaran.", "Admin menyusun jadwal serta menyelesaikan konflik yang ditolak sistem.", "Guru menginput nilai per pengajaran dan bulan.", "Admin/Kepala Sekolah memonitor kelengkapan nilai dan laporan.", "Admin menonaktifkan data lama tanpa menghapus histori."]),
        h1("6. SOP Deployment"), numbered(["Siapkan server PHP 8.3+, database, Composer, Node.js, dan HTTPS.", "Pasang source code dan arahkan document root ke public/.", "Isi .env production tanpa membagikan credential melalui dokumen umum.", "Jalankan composer install --no-dev, npm run build, migrate --force, storage:link, dan optimize.", "Atur write permission hanya untuk storage/ dan bootstrap/cache/.", "Uji login, upload logo, jadwal, nilai, laporan, dan PDF.", "Aktifkan backup database dan monitoring log."]),
        h1("7. Backup dan Pemulihan"), table(["Objek", "Frekuensi", "Metode"], [["Database", "Harian atau sebelum perubahan besar", "mysqldump/fitur backup hosting; simpan terenkripsi."], ["storage/app/public", "Mingguan atau setelah perubahan logo/file", "Salin ke lokasi backup terpisah."], ["Source code/config template", "Setiap rilis", "Repository/artifact; jangan commit .env."], ["Dokumen", "Setiap finalisasi", "Simpan PDF/DOCX dan salinan offline."]]),
        h1("8. Troubleshooting"), table(["Gejala", "Langkah"], [["HTTP 500", "Periksa storage/logs/laravel.log, APP_DEBUG tetap false di production, dan permission folder."], ["Database gagal", "Periksa DB_HOST, DB_PORT, credential, service database, migration."], ["403", "Periksa role/permission dan status akun."], ["419", "Periksa session, cookie, HTTPS, dan APP_URL."], ["Aset tidak tampil", "Jalankan build dan periksa manifest."], ["PDF gagal", "Periksa data/logo, permission, memori PHP, dan log."], ["Jadwal/nilai ditolak", "Baca pesan validasi dan periksa periode, pengajaran, kelas, serta hak pengguna."]]),
        h1("9. Pengelolaan Insiden dan Perubahan"), bullets(["Catat waktu, pengguna, langkah, pesan error, URL, dan screenshot.", "Bedakan bug dalam scope, kesalahan data, masalah infrastruktur, dan permintaan fitur baru.", "Jangan mengubah database langsung tanpa backup dan persetujuan.", "Uji perbaikan pada lingkungan non-production sebelum diterapkan.", "Permintaan fitur baru dicatat sebagai change request dengan dampak biaya/waktu."]),
        h1("10. Form Catatan Pendampingan"), table(["Tanggal", "Topik", "Peserta", "Masalah/Pertanyaan", "Tindak Lanjut"], [["", "", "", "", ""], ["", "", "", "", ""], ["", "", "", "", ""], ["", "", "", "", ""]]),
        h1("11. Checklist Serah Pengetahuan"), bullets(["Admin dapat melakukan setup periode dan master data.", "Guru dapat menginput nilai dan membuka laporan.", "Siswa/Kepala Sekolah memahami batas akses.", "Tim teknis dapat menjalankan instalasi dan backup dasar.", "Dokumen ERD, flow, testing, sistem, demo, dan pendampingan telah diterima.", "Credential produksi telah diserahkan melalui media aman."]),
    ])
    return [erd, useflow, blackbox, documentation, demo, support]


def copy_supporting_files():
    supporting = OUTPUT / "Lampiran-Pendukung"
    supporting.mkdir(parents=True, exist_ok=True)
    for source in ["docs/erd.dbml", "docs/roles-permissions.md", "docs/system-flow.md", "docs/deployment.md", "docs/black-box-testing.md", "docs/buku-panduan-penggunaan.pdf"]:
        path = ROOT / source
        if path.exists():
            shutil.copy2(path, supporting / path.name)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "erd": generate_erd(),
        "usecase": generate_use_case(),
        "login": generate_flow("login", "FLOW LOGIN DAN AKSES", [("Buka Login", "Pengguna memasukkan username dan password"), ("Validasi", "Kredensial diperiksa",), ("Akun Aktif?", "Status user harus aktif"), ("Muat Role", "Permission dan profil dimuat"), ("Dashboard", "Menu dan data sesuai role")], {2}),
        "setup": generate_flow("setup", "FLOW SETUP AKADEMIK", [("Aktifkan Periode", "Tahun ajaran dan semester"), ("Isi Master Data", "Guru, siswa, kelas, mapel, jam"), ("Tempatkan Siswa", "Anggota kelas dan histori"), ("Buat Pengajaran", "Guru-mapel-kelas-semester"), ("Susun Jadwal", "Hari dan jam pelajaran"), ("Siap Operasional", "Guru dapat mengisi nilai")]),
        "schedule": generate_flow("jadwal", "FLOW PENYUSUNAN JADWAL", [("Pilih Periode", "Tahun ajaran, semester, kelas"), ("Pilih Pengajaran", "Guru dan mata pelajaran"), ("Pilih Hari/Jam", "Slot pelajaran"), ("Cek Bentrok?", "Kelas, guru, dan duplikasi"), ("Simpan", "Jika valid"), ("Tampilkan", "Jadwal sesuai filter")], {3}),
        "grade": generate_flow("nilai", "FLOW INPUT NILAI", [("Pilih Pengajaran", "Guru dan kelas"), ("Pilih Bulan", "1 sampai 12"), ("Muat Siswa", "Anggota kelas aktif"), ("Isi M1–M4", "0–100 atau kosong"), ("Validasi", "Hak guru, anggota, rentang"), ("Simpan Massal", "Upsert dalam transaksi")], {4}),
        "report": generate_flow("laporan", "FLOW LAPORAN", [("Pilih Jenis", "Jadwal atau nilai"), ("Terapkan Filter", "Periode, kelas, guru, mapel"), ("Batasi Sesuai Role", "Guru/siswa otomatis dibatasi"), ("Preview Web", "Data dan identitas sekolah"), ("Print / PDF", "A4 portrait/landscape")]),
        "architecture": generate_architecture(),
    }
    specs = build_specs(diagrams)
    for spec in specs:
        render_docx(spec)
        render_html(spec)
    copy_supporting_files()
    manifest = "DOKUMEN PAKET COMPLETE SISDAR\n\n" + "\n".join(f"{index}. {spec.title} — {spec.filename}.pdf / .docx" for index, spec in enumerate(specs, 1))
    manifest += f"\n\nTanggal penyusunan: {DATE}\nInstansi: {SCHOOL}\n"
    (OUTPUT / "00-DAFTAR-DOKUMEN.txt").write_text(manifest, encoding="utf-8")
    print(f"Dokumen sumber dibuat di {OUTPUT}")


if __name__ == "__main__":
    main()
